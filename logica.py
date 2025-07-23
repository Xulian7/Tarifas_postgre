import pandas as pd
import tkinter as tk
import tkinter.font as tkFont
from tkinter import filedialog, messagebox, ttk
import os
from datetime import datetime, timedelta, time
from tkcalendar import DateEntry
import locale
from tkinter import filedialog
import math
from PIL import ImageGrab
import io
import win32clipboard
import json
from openpyxl import Workbook
import ctypes
from tkinter import Toplevel
from sqlalchemy import select, insert, update, delete, case, func, text, and_, or_, Table, MetaData
from sqlalchemy.exc import SQLAlchemyError
from conexion import engine, registros
from sqlalchemy.orm import Session
from conexion import engine
from conexion import clientes as tabla_clientes, registros as tabla_registros, propietario as tabla_propietario
from conexion import cuentas as tabla_cuentas, otras_deudas as tabla_otras_deudas
from docx import Document
from docx2pdf import convert
from num2words import num2words
import tempfile
import re

JSON_PATH = 'diccionarios/columnas.json'
XLSX_PATH = 'diccionarios/estructura.xlsx'
total_deuda_float = 0.0
total_abonos_float = 0.0
# Establecer configuraciones locales - español
locale.setlocale(locale.LC_ALL, 'es_CO.utf8')
ventana_clientes = None  # Variable global dentro del módulo
ventana_atrasos = None  # Variable global para evitar duplicados

# ---------- Cargar datos desde PostgreSQL ----------
def cargar_db(tree, entry_cedula, entry_nombre, entry_placa, entry_referencia, entry_fecha, combo_tipo, combo_nequi, combo_verificada):
    try:
        cedula     = entry_cedula.get()
        nombre     = entry_nombre.get()
        placa      = entry_placa.get()
        referencia = entry_referencia.get()
        fecha      = entry_fecha.get()
        tipo       = combo_tipo.get()
        nequi      = combo_nequi.get()
        verificada = combo_verificada.get()

        if fecha:
            fecha = datetime.strptime(fecha, "%d-%m-%Y").strftime("%Y-%m-%d")

        condiciones = []

        if cedula:
            condiciones.append(tabla_registros.c.cedula == cedula)
        if nombre:
            condiciones.append(tabla_registros.c.nombre.ilike(f"%{nombre}%"))
        if placa:
            condiciones.append(tabla_registros.c.placa.ilike(f"%{placa}%"))
        if referencia:
            condiciones.append(tabla_registros.c.referencia.ilike(f"%{referencia}%"))
        if fecha:
            condiciones.append(tabla_registros.c.fecha_registro == fecha)
        if tipo:
            condiciones.append(tabla_registros.c.tipo == tipo)
        if nequi:
            condiciones.append(tabla_registros.c.nombre_cuenta == nequi)
        if verificada:
            condiciones.append(tabla_registros.c.verificada == verificada)

        stmt = (
            select(
                tabla_registros.c.id,
                tabla_registros.c.fecha_sistema,
                tabla_registros.c.fecha_registro,
                tabla_registros.c.cedula,
                tabla_registros.c.nombre,
                tabla_registros.c.placa,
                tabla_registros.c.valor,
                tabla_registros.c.saldos,
                tabla_registros.c.motivo,
                tabla_registros.c.tipo,
                tabla_registros.c.nombre_cuenta,
                tabla_registros.c.referencia,
                tabla_registros.c.verificada
            )
            .select_from(tabla_registros.outerjoin(tabla_propietario, tabla_registros.c.placa == tabla_propietario.c.placa))
            .where(and_(*condiciones))
        )

        with engine.connect() as conn:
            rows = conn.execute(stmt).fetchall()

        tree.delete(*tree.get_children())

        rows.sort(key=lambda x: (str(x[3]), str(x[1])))

        for row in rows:
            fecha_sistema  = pd.to_datetime(row[1]).strftime('%d-%m-%Y')
            fecha_registro = pd.to_datetime(row[2]).strftime('%d-%m-%Y')

            values = list(row)
            values[1] = fecha_sistema
            values[2] = fecha_registro

            tree.insert("", "end", values=values)

        for col in tree["columns"]:
            max_width = max([tkFont.Font().measure(col)] + [tkFont.Font().measure(str(value)) for value in rows])
            tree.column(col, width=max_width)

    except Exception as e:
        messagebox.showerror("Error", f"No se pudo cargar los datos desde PostgreSQL: {e}")

# ---------- Agregar registro a la base de datos ----------
def agregar_registro(tree, entry_hoy, entry_cedula, entry_nombre, entry_placa, entry_monto, entry_saldos, combo_motivo,
                     entry_referencia, entry_fecha, combo_tipo, combo_nequi, combo_verificada,
                     listbox_sugerencias):

    # Obtener valores
    cedula = entry_cedula.get().strip()
    nombre = entry_nombre.get().strip()
    placa = entry_placa.get().strip()
    valor = entry_monto.get().strip()
    saldos = entry_saldos.get().strip()
    referencia = entry_referencia.get().strip()
    fecha_hoy = entry_hoy.get().strip()
    fecha = entry_fecha.get().strip()
    tipo = combo_tipo.get().strip()
    nequi = combo_nequi.get().strip()
    verificada = "No"
    motivo = combo_motivo.get().strip()

    # Validación de campos obligatorios
    campos_faltantes = []
    if not cedula: campos_faltantes.append("Cédula")
    if not nombre: campos_faltantes.append("Nombre")
    if not placa: campos_faltantes.append("Placa")
    if not valor: campos_faltantes.append("Valor")
    if not saldos: campos_faltantes.append("Saldos")
    if not fecha_hoy: campos_faltantes.append("Fecha de hoy")
    if not tipo: campos_faltantes.append("Tipo")
    if tipo.lower() not in ["efectivo", "ajuste p/p"]:
        if not referencia: campos_faltantes.append("Referencia")
        if not nequi: campos_faltantes.append("Nequi")

    if campos_faltantes:
        mensaje_error = "Faltan valores obligatorios:\n- " + "\n- ".join(campos_faltantes)
        messagebox.showerror("Error", mensaje_error)
        return

    try:
        valor = float(valor)
        saldos = float(saldos) if saldos else 0.0

        if saldos != 0 and motivo.lower() == "n-a":
            messagebox.showwarning("Motivo requerido", "Debe asignar un motivo para el valor ingresado.")
            return
    except ValueError:
        messagebox.showerror("Error", "Valor y saldos deben ser numéricos.")
        return


    fecha_hoy_bd = convertir_fecha(fecha_hoy)
    if not fecha_hoy_bd:
        return

    # Si tipo es "efectivo", se fuerza la fecha_registro = fecha_hoy
    if tipo.lower() == "efectivo":
        fecha_bd = fecha_hoy_bd
    else:
        fecha_bd = convertir_fecha(fecha)
        if not fecha_bd:
            return
        
        # Validar que la fecha no sea futura
    if fecha_bd > datetime.today().date():
        messagebox.showwarning("Fecha inválida", "No puedes registrar fechas posteriores al día de hoy.")
        return


    try:
        with engine.begin() as conn:
            # Validar referencia duplicada
            if referencia:
                ref_check = select(
                    tabla_registros.c.referencia,
                    tabla_registros.c.nombre,
                    tabla_registros.c.fecha_registro,
                    tabla_registros.c.id
                ).where(tabla_registros.c.referencia == referencia)

                resultado = conn.execute(ref_check).first()

                if resultado:
                    referencia_existente = resultado.referencia
                    nombre_existente = resultado.nombre
                    fecha_existente = resultado.fecha_registro
                    id_existente = resultado.id

                    # Formatea fecha si es datetime
                    if hasattr(fecha_existente, "strftime"):
                        fecha_formateada = fecha_existente.strftime("%d/%m/%Y %H:%M")
                    else:
                        fecha_formateada = str(fecha_existente)

                    messagebox.showwarning(
                        "⚠️ Referencia duplicada",
                        f"La referencia '{referencia_existente}' ya existe.\n\n"
                        f"🔹 **Nombre:** {nombre_existente}\n"
                        f"📅 **Fecha de registro:** {fecha_formateada}\n"
                        f"🆔 **ID transacción:** {id_existente}"
                    )


                    return


            # Validar combinación única
            check_cliente = select(tabla_clientes.c.cedula).where(
                and_(
                    tabla_clientes.c.cedula == cedula,
                    tabla_clientes.c.nombre == nombre,
                    tabla_clientes.c.placa == placa
                )
            )
            if conn.execute(check_cliente).rowcount != 1:
                messagebox.showerror("Error", "La combinación de cédula, nombre y placa no es única o no existe.")
                return

            # Confirmar inserción
            confirmar = messagebox.askyesno("Confirmar", "¿Deseas grabar este registro?")
            if not confirmar:
                messagebox.showinfo("Cancelado", "La operación fue cancelada.")
                return

            # Insertar el registro
            insertar = insert(tabla_registros).values(
                fecha_sistema=fecha_hoy_bd,
                fecha_registro=fecha_bd,
                cedula=cedula,
                nombre=nombre,
                placa=placa,
                valor=valor,
                saldos=saldos,
                motivo=motivo,
                tipo=tipo,
                nombre_cuenta=nequi,
                referencia=referencia,
                verificada=verificada
            )
            conn.execute(insertar)

        mostrar_msgbox_exito(
            entry_cedula,
            lambda: limpiar_formulario(entry_cedula, entry_nombre, entry_placa, entry_monto, entry_saldos,
                                       combo_motivo, entry_referencia, entry_fecha, combo_tipo,
                                       combo_nequi, combo_verificada, listbox_sugerencias, tree),
            lambda: limpiar_parcial(entry_monto, entry_saldos, combo_motivo, entry_referencia,
                                    entry_fecha, combo_tipo, combo_nequi, combo_verificada,
                                    listbox_sugerencias, tree)
        )
        
        hoy = datetime.today().strftime("%d/%m/%Y")
        entry_fecha.delete(0, tk.END)
        entry_fecha.insert(0, hoy)
        
    except Exception as e:
        messagebox.showerror("Error", f"Error en base de datos:\n{e}")

# ---------- Mostrar mensaje de éxito en un MsgBox ----------
def mostrar_msgbox_exito(entry_cedula, limpiar_funcion, parcial_funcion):
    ventana = Toplevel()
    ventana.title("Éxito")
    ventana.geometry("350x150")
    ventana.resizable(False, False)
    ventana.grab_set()  # Ventana modal
    label = tk.Label(ventana, text="Registro agregado correctamente.", font=("Arial", 12))
    label.pack(pady=20)
    botones_frame = tk.Frame(ventana)
    botones_frame.pack(pady=10)
    
    btn_generar = tk.Button(botones_frame, text="Generar Recibo", width=15,
    command=lambda: [
        mostrar_registros(entry_cedula),
        limpiar_funcion(),
        ventana.destroy()
    ])
    btn_generar.pack(side="left", padx=10)
    
    btn_aceptar = tk.Button(botones_frame, text="Repetir Cliente", width=15,
    command=lambda: [
        parcial_funcion(),
        ventana.destroy()])
    btn_aceptar.pack(side="right", padx=10)

# ---------- Limpiar formulario completo ----------
def limpiar_formulario(entry_cedula, entry_nombre, entry_placa, entry_monto, entry_saldos, combo_motivo, entry_referencia, entry_fecha,
combo_tipo, combo_nequi, combo_verificada, listbox_sugerencias, tree):
    # Limpiar campos de texto (Entry)
    entry_cedula.focus_set()
    entry_cedula.delete(0, tk.END)
    entry_nombre.delete(0, tk.END)
    entry_placa.delete(0, tk.END)
    entry_monto.delete(0, tk.END)
    entry_saldos.delete(0, tk.END)
    entry_referencia.delete(0, tk.END)
    entry_fecha.delete(0, tk.END)
    combo_motivo.set('N-a')  # Resetear el ComboBox de Motivo
    
    
    # Limpiar los Combobox
    combo_tipo.set('')  # Resetear el ComboBox de Tipo
    combo_nequi.set('')  # Resetear el ComboBox de Nequi
    combo_verificada.set('No')  # Resetear el ComboBox de Verificada
    listbox_sugerencias.grid_forget()
    
    # Limpiar Treeview
    for row in tree.get_children():
        tree.delete(row)

# ---------- Limpiar formulario parcial ----------       
def limpiar_parcial(entry_monto, entry_saldos, combo_motivo, entry_referencia, entry_fecha,
combo_tipo, combo_nequi, combo_verificada, listbox_sugerencias, tree):
    # Limpiar campos de texto (Entry)
    entry_monto.delete(0, tk.END)
    entry_saldos.delete(0, tk.END)
    entry_referencia.delete(0, tk.END)
    entry_fecha.delete(0, tk.END)
    combo_motivo.set('N-a')  # Resetear el ComboBox de Motivo
    # Limpiar los Combobox
    combo_tipo.set('')  # Resetear el ComboBox de Tipo
    combo_nequi.set('')  # Resetear el ComboBox de Nequi
    combo_verificada.set('No')  # Resetear el ComboBox de Verificada
    listbox_sugerencias.grid_forget()
    # Limpiar Treeview
    for row in tree.get_children():
        tree.delete(row)

# ---------- Cargar opciones de Nequi desde la base de datos ----------
def cargar_nequi_opciones():
    try:
        with engine.connect() as conn:
            stmt = select(tabla_cuentas.c.nombre_cuenta)
            rows = conn.execute(stmt).fetchall()

            return [row[0] for row in rows]

    except Exception as e:
        print(f"💥 Error al cargar los datos de Nequi: {e}")
        return []

# ---------- Convertir fecha de string a objeto date ----------
def convertir_fecha(fecha_str):
    """Convierte una fecha en formato dd-mm-yyyy a un objeto date."""
    try:
        return datetime.strptime(fecha_str, "%d-%m-%Y").date()
    except ValueError:
        messagebox.showerror("Error", "Formato de fecha incorrecto. Use dd-mm-yyyy.")
        return None

# ---------- Ajustar automáticamente el ancho de las columnas ----------
def ajustar_columnas(tree):
    """Ajusta automáticamente el ancho de las columnas en función del contenido."""
    for col in tree["columns"]:
        tree.column(col, anchor="center")  # Justificar contenido al centro
        max_len = len(col)  # Inicia con el ancho del encabezado
        for item in tree.get_children():
            text = str(tree.item(item, "values")[tree["columns"].index(col)])
            max_len = max(max_len, len(text))
        tree.column(col, width=max_len * 10)  # Ajusta el ancho en función del contenido

# ---------- Obtener datos de clientes desde PostgreSQL ----------
def obtener_datos_clientes():
    """Obtiene los datos de la tabla clientes desde la base de datos PostgreSQL y formatea Capital y Fecha_inicio."""
    try:
        stmt = (
            select(
                tabla_clientes.c.cedula,
                tabla_clientes.c.nombre,
                tabla_clientes.c.nacionalidad,
                tabla_clientes.c.telefono,
                tabla_clientes.c.direccion,
                tabla_clientes.c.placa,
                tabla_propietario.c.modelo,
                tabla_propietario.c.tarjeta_propiedad,
                tabla_clientes.c.fecha_inicio,
                tabla_clientes.c.fecha_final,
                tabla_clientes.c.tipo_contrato,
                tabla_clientes.c.valor_cuota,
                tabla_clientes.c.estado,
                tabla_clientes.c.otras_deudas,
                tabla_clientes.c.visitador,
                tabla_clientes.c.referencia,
                tabla_clientes.c.telefono_ref
            )
            .select_from(
                tabla_clientes.outerjoin(
                    tabla_propietario, tabla_clientes.c.placa == tabla_propietario.c.placa
                )
            )
        )

        with engine.connect() as conn:
            result = conn.execute(stmt).fetchall()

        datos_formateados = []
        for fila in result:
            (
                cedula, nombre, nacionalidad, telefono, direccion, 
                placa, modelo, tarjeta_propiedad, fecha_inicio, fecha_final, 
                tipo_contrato, valor_cuota, estado, otras_deudas, 
                visitador, referencia, telefono_ref
            ) = fila

            if fecha_inicio:
                try:
                    fecha_inicio = fecha_inicio.strftime("%d-%m-%Y")
                except Exception:
                    fecha_inicio = "Formato Inválido"

            datos_formateados.append((
                cedula, nombre, nacionalidad, telefono, direccion, 
                placa, modelo, tarjeta_propiedad, fecha_inicio, fecha_final, 
                tipo_contrato, valor_cuota, estado, otras_deudas, 
                visitador, referencia, telefono_ref
            ))

        return datos_formateados

    except Exception as e:
        print(f"Error al obtener datos de clientes: {e}")
        return []

# ---------- Abrir ventana de clientes ----------
def abrir_ventana_clientes():
    
    global ventana_clientes

    if ventana_clientes and ventana_clientes.winfo_exists():
        ventana_clientes.lift()  # Trae la ventana al frente si ya existe
        return
    
    ventana_clientes = tk.Toplevel()
    ventana_clientes.title("Clientes")
    ventana_clientes.geometry("900x600")
    icono_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'img', 'inicio.ico')
    if os.path.exists(icono_path):
        ventana_clientes.iconbitmap(icono_path)
    else:
        print("No se encontró el icono en la ruta especificada")

    # Crear un Frame para contener el Treeview y la Scrollbar
    frame_tree = ttk.Frame(ventana_clientes)
    frame_tree.grid(row=0, column=0, columnspan=4, sticky="nsew", padx=10, pady=10)

    # Crear el Treeview dentro del Frame
    columnas = ("Cédula", "Nombre","Nacionalidad", "Teléfono", "Dirección", 
                "Placa","Modelo","Tarjeta propiedad", "Fecha Inicio", "Fecha Final", "Tipo Contrato", "Valor Cuota", "Estado", "Total_inicial", "Visitador", "Referencia", "Telefono_ref")

    tree = ttk.Treeview(frame_tree, columns=columnas, show="headings")

    # Configurar encabezados y justificar contenido al centro
    for col in columnas:
        tree.heading(col, text=col)
        tree.column(col, anchor="center")

    # Crear Scrollbar dentro del Frame
    scrollbar = ttk.Scrollbar(frame_tree, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=scrollbar.set)

    # Ubicar Treeview y Scrollbar con grid dentro del Frame
    tree.grid(row=0, column=0, sticky="nsew")
    scrollbar.grid(row=0, column=1, sticky="ns")

    # Configurar expansión del Frame
    frame_tree.columnconfigure(0, weight=1)
    frame_tree.rowconfigure(0, weight=1)

    # Llenar el Treeview con datos de la base de datos
    for fila in obtener_datos_clientes():
        tree.insert("", "end", values=fila)
    # Ajustar automáticamente el ancho de las columnas después de insertar los datos
    ajustar_columnas(tree)
    
    global datos_originales
    datos_originales = [tree.item(item)["values"] for item in tree.get_children()]
    
    # Función para configurar correctamente los DateEntry
    def create_date_entry(parent):
        return DateEntry(parent, width=27, background='darkblue', 
                        foreground='white', borderwidth=2, 
                        date_pattern='dd-MM-yyyy',  # Establecer el formato Día-Mes-Año
                        locale='es_ES')

    # Frame para los Labels y Entries
    frame_form = ttk.LabelFrame(ventana_clientes, text="Información del Cliente")
    frame_form.grid(row=1, column=0, columnspan=6, padx=10, pady=10, sticky="nsew")

    # Diccionario para almacenar las entradas
    entries = {}
    
    cedula_var = tk.StringVar()
    # Función para validar que solo sean números (Si son letras las borra)
    def validar_cedula(*args):
        valor = cedula_var.get()
        if not valor.isdigit():
            cedula_var.set("".join(filter(str.isdigit, valor)))  # Elimina caracteres no numéricos

    # Agregando manualmente cada etiqueta y entrada en un diseño de 3 columnas
    cedula_var.trace_add("write", validar_cedula)
    lbl_cedula = ttk.Label(frame_form, text="Cédula:")
    lbl_cedula.grid(row=0, column=0, padx=5, pady=5, sticky="w")
    entries["Cédula"] = ttk.Entry(frame_form, textvariable=cedula_var, width=30)
    entries["Cédula"].grid(row=0, column=1, padx=5, pady=5, sticky="w")

    nombre_var = tk.StringVar()
    nombre_var.trace_add("write", lambda *args: nombre_var.set(nombre_var.get().title()))
    nombre_var.trace_add("write", lambda *args: filtrar_treeview(tree, nombre_var))
    lbl_nombre = ttk.Label(frame_form, text="Nombre:")
    lbl_nombre.grid(row=0, column=2, padx=5, pady=5, sticky="w")
    
    def filtrar_treeview(tree, nombre_var):
        filtro = nombre_var.get().strip().lower()
        items = tree.get_children()

        if not filtro:
            tree.delete(*items)
            for fila in datos_originales:
                tree.insert("", "end", values=fila)
            return

        tree.delete(*items)
        for fila in datos_originales:
            if filtro in fila[1].lower():
                tree.insert("", "end", values=fila)
    
    # Función para consultar datos del vehículo
    def consultar_datos_vehiculo(*args):
        placa = placa_var.get()
        if not placa:
            return

        try:
            stmt = select(
                tabla_propietario.c.modelo,
                tabla_propietario.c.tarjeta_propiedad
            ).where(tabla_propietario.c.placa == placa)

            with engine.connect() as conn:
                resultado = conn.execute(stmt).fetchone()

            if resultado:
                modelo, tarjeta = resultado

                entries["Modelo"].config(state="normal")
                entries["Modelo"].delete(0, tk.END)
                entries["Modelo"].insert(0, modelo)
                entries["Modelo"].config(state="readonly")

                entries["Tarjeta_propiedad"].config(state="normal")
                entries["Tarjeta_propiedad"].delete(0, tk.END)
                entries["Tarjeta_propiedad"].insert(0, tarjeta)
                entries["Tarjeta_propiedad"].config(state="readonly")

        except Exception as e:
            print(f"Error al consultar el vehículo: {e}")


    entries["Nombre"] = ttk.Entry(frame_form, textvariable=nombre_var, width=30)
    entries["Nombre"].grid(row=0, column=3, padx=5, pady=5, sticky="w")
    nacion_var = tk.StringVar()
    nacion_var.trace_add("write", lambda *args: nacion_var.set(nacion_var.get().title()))
    lbl_nacion = ttk.Label(frame_form, text="Nacionalidad:")
    lbl_nacion.grid(row=0, column=4, padx=4, pady=5, sticky="w")
    entries["Nacionalidad"] = ttk.Entry(frame_form, textvariable=nacion_var, width=30)
    entries["Nacionalidad"].grid(row=0, column=5, padx=5, pady=5, sticky="w")
    lbl_telefono = ttk.Label(frame_form, text="Teléfono:")
    lbl_telefono.grid(row=1, column=0, padx=0, pady=5, sticky="w")
    entries["Teléfono"] = ttk.Entry(frame_form, width=30)
    entries["Teléfono"].grid(row=1, column=1, padx=5, pady=5, sticky="w")
    lbl_direccion = ttk.Label(frame_form, text="Dirección:")
    lbl_direccion.grid(row=1, column=2, padx=5, pady=5, sticky="w")
    entries["Dirección"] = ttk.Entry(frame_form, width=30)
    entries["Dirección"].grid(row=1, column=3, padx=5, pady=5, sticky="w")
    placa_var = tk.StringVar()
    placa_var.trace_add("write", lambda *args: placa_var.set(placa_var.get().upper()))
    placa_var.trace_add("write", consultar_datos_vehiculo)  # Consultar modelo y tarjeta
    lbl_placa = ttk.Label(frame_form, text="Placa:")
    lbl_placa.grid(row=1, column=4, padx=5, pady=5, sticky="w")
    entries["Placa"] = ttk.Entry(frame_form, textvariable=placa_var, width=30)
    entries["Placa"].grid(row=1, column=5, padx=5, pady=5, sticky="w")
    lbl_modelo = ttk.Label(frame_form, text="Modelo:")
    lbl_modelo.grid(row=2, column=0, padx=5, pady=5, sticky="w")
    entries["Modelo"] = ttk.Entry(frame_form, width=30, state="readonly")
    entries["Modelo"].grid(row=2, column=1, padx=5, pady=5, sticky="w")
    lbl_tarjeta_propiedad = ttk.Label(frame_form, text="Tarjeta Propiedad:")
    lbl_tarjeta_propiedad.grid(row=2, column=2, padx=5, pady=5, sticky="w")
    entries["Tarjeta_propiedad"] = ttk.Entry(frame_form, width=30, state="readonly")
    entries["Tarjeta_propiedad"].grid(row=2, column=3, padx=5, pady=5, sticky="w")
    lbl_fecha_inicio = ttk.Label(frame_form, text="Fecha Inicio:")
    lbl_fecha_inicio.grid(row=2, column=4, padx=5, pady=5, sticky="w")
    entries["Fecha Inicio"] = create_date_entry(frame_form)
    entries["Fecha Inicio"].grid(row=2, column=5, padx=5, pady=5, sticky="w")
    lbl_fecha_final = ttk.Label(frame_form, text="Fecha Final:")
    lbl_fecha_final.grid(row=3, column=0, padx=5, pady=5, sticky="w")
    entries["Fecha Final"] = ttk.Entry(frame_form, width=30)
    entries["Fecha Final"].grid(row=3, column=1, padx=5, pady=5, sticky="w")
    lbl_tipo_contrato = ttk.Label(frame_form, text="Tipo Contrato:")
    lbl_tipo_contrato.grid(row=3, column=2, padx=5, pady=5, sticky="w")
    tipos_opciones = ["Dia", "Semana", "Quincena", "Mes", "Sin asignar"]
    entries["Tipo Contrato"] = ttk.Combobox(frame_form, values=tipos_opciones, state="readonly", width=30)
    entries["Tipo Contrato"].grid(row=3, column=3, padx=5, pady=5, sticky="w")
    lbl_valor_cuota = ttk.Label(frame_form, text="Valor Cuota:")
    lbl_valor_cuota.grid(row=3, column=4, padx=5, pady=5, sticky="w")
    entries["Valor Cuota"] = ttk.Entry(frame_form, width=30)
    entries["Valor Cuota"].grid(row=3, column=5, padx=5, pady=5, sticky="w")
    lbl_estado = ttk.Label(frame_form, text="Estado:")
    lbl_estado.grid(row=0, column=6, padx=5, pady=5, sticky="w")
    estado_opciones = ["","activo", "inactivo"]
    combo_estado = ttk.Combobox(frame_form, values=estado_opciones, width=30)
    combo_estado.grid(row=0, column=7, padx=5, pady=5, sticky="w")
    lbl_otras_deudas = ttk.Label(frame_form, text="Otras deudas:")
    lbl_otras_deudas.grid(row=1, column=6, padx=5, pady=5, sticky="w")
    entries["Otras deudas"] = ttk.Entry(frame_form, width=30)
    entries["Otras deudas"].grid(row=1, column=7, padx=5, pady=5, sticky="w")
    lbl_visitador = ttk.Label(frame_form, text="Visitador:")
    lbl_visitador.grid(row=2, column=6, padx=5, pady=5, sticky="w")
    entries["Visitador"] = ttk.Entry(frame_form, width=30)
    entries["Visitador"].grid(row=2, column=7, padx=5, pady=5, sticky="w")
    lbl_referencia = ttk.Label(frame_form, text="Referencia:")
    lbl_referencia.grid(row=3, column=6, padx=5, pady=5, sticky="w")
    entries["Referencia"] = ttk.Entry(frame_form, width=30)
    entries["Referencia"].grid(row=3, column=7, padx=5, pady=5, sticky="w")
    lbl_tel_referencia = ttk.Label(frame_form, text="Telefono Ref:")
    lbl_tel_referencia.grid(row=3, column=8, padx=5, pady=5, sticky="w")
    entries["Telefono Ref"] = ttk.Entry(frame_form, width=30)
    entries["Telefono Ref"].grid(row=3, column=9, padx=5, pady=5, sticky="w")


    def limpiar_formulario():
        """Limpia todos los campos de entrada en el formulario."""
        for entry in entries.values():
                entry.delete(0, "end")
        
        entries["Tipo Contrato"].set("")  # Resetear el Combobox
        entries["Tarjeta_propiedad"].config(state="normal")
        entries["Tarjeta_propiedad"].delete(0, tk.END)
        entries["Tarjeta_propiedad"].config(state="readonly")
        combo_estado.set("")
        
    def convertir_fecha_formato_sqlite(fecha_ui):
        """Convierte una fecha de formato dd-mm-yyyy a yyyy-mm-dd"""
        try:
            dia, mes, año = fecha_ui.split('-')
            return f"{año}-{mes}-{dia}"
        except ValueError:
            messagebox.showerror("Error", f"Formato de fecha inválido: {fecha_ui}")
            return None

    def generar_contrato(valores_dict, plantilla_path="diccionarios/contrato.docx"):
        if not os.path.exists(plantilla_path):
            return None  # ⛔ No hay plantilla

        # Cargar plantilla
        doc = Document(plantilla_path)

        # Reemplazar en párrafos
        for p in doc.paragraphs:
            for key, val in valores_dict.items():
                p.text = p.text.replace(f"{{{key}}}", str(val))

        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in valores_dict.items():
                        cell.text = cell.text.replace(f"{{{key}}}", str(val))

        # Guardar archivo temporal
        temp_docx = os.path.join(tempfile.gettempdir(), "contrato_temp.docx")
        doc.save(temp_docx)

        # Ruta: Escritorio/contratos
        escritorio = os.path.join(os.path.expanduser("~"), "Desktop")
        ruta_contratos = os.path.join(escritorio, "contratos")
        os.makedirs(ruta_contratos, exist_ok=True)

        # Nombre del archivo de salida
        nombre = valores_dict["NOMBRE"].replace(" ", "")
        placa = valores_dict["PLACA"].replace(" ", "").upper()
        salida_pdf = os.path.join(ruta_contratos, f"Contrato {nombre} {placa}.pdf")

        # Convertir a PDF
        convert(temp_docx, salida_pdf)
        return salida_pdf


    def registrar_cliente():
        # Obtener valores de los campos
        valores = [
            entries["Cédula"].get().strip(),
            entries["Nombre"].get().strip(),
            entries["Nacionalidad"].get().strip(),
            entries["Teléfono"].get().strip(),
            entries["Dirección"].get().strip(),
            entries["Placa"].get().strip(),
            convertir_fecha_formato_sqlite(entries["Fecha Inicio"].get().strip()),
            entries["Fecha Final"].get().strip(),
            entries["Tipo Contrato"].get().strip(),
            entries["Valor Cuota"].get().strip(),
            combo_estado.get().strip(),
            entries["Otras deudas"].get().strip(),
            entries["Visitador"].get().strip(),
            entries["Referencia"].get().strip(),
            entries["Telefono Ref"].get().strip()
        ]

        if '' in valores:
            messagebox.showerror("Error", "Todos los campos deben estar llenos.")
            ventana_clientes.focus_force()
            return

        try:
            with engine.begin() as conn:
                stmt_prop = select(tabla_propietario.c.placa).where(tabla_propietario.c.placa == valores[5]).limit(1)
                resultado = conn.execute(stmt_prop).fetchone()
                if not resultado:
                    messagebox.showwarning("Advertencia", f"La Placa {valores[5]} no está registrada en la base de datos de propietarios.")
                    return

                stmt_check = select(
                    tabla_clientes.c.cedula,
                    tabla_clientes.c.placa
                ).where(
                    (tabla_clientes.c.cedula == valores[0]) | (tabla_clientes.c.placa == valores[5])
                ).limit(1)
                resultado = conn.execute(stmt_check).fetchone()
                if resultado:
                    mensaje = "No se puede registrar el cliente porque:\n"
                    if resultado[0] == valores[0]:
                        mensaje += f"- La Cédula {resultado[0]} ya está registrada.\n"
                    if resultado[1] == valores[5]:
                        mensaje += f"- La Placa {resultado[1]} ya está asignada a otro cliente.\n"
                    messagebox.showwarning("Advertencia", mensaje)
                    return

                # Insertar cliente
                conn.execute(insert(tabla_clientes), {
                    "cedula": valores[0],
                    "nombre": valores[1],
                    "nacionalidad": valores[2],
                    "telefono": valores[3],
                    "direccion": valores[4],
                    "placa": valores[5],
                    "fecha_inicio": valores[6],
                    "fecha_final": valores[7],
                    "tipo_contrato": valores[8],
                    "valor_cuota": valores[9],
                    "estado": valores[10],
                    "otras_deudas": valores[11],
                    "visitador": valores[12],
                    "referencia": valores[13],
                    "telefono_ref": valores[14]
                })

                conn.execute(insert(tabla_otras_deudas), {
                    "cedula": valores[0],
                    "placa": valores[5],
                    "fecha_deuda": valores[6],
                    "descripcion": "Cuota Inicial",
                    "valor": valores[11]
                })

            # Preparar valores para contrato
            fecha_dt = datetime.strptime(valores[6], "%Y-%m-%d")
            fecha_dia_siguiente = fecha_dt + timedelta(days=1)
            valor_cuota = valores[9]
            valores_dict = {
                "CEDULA": valores[0],
                "NOMBRE": valores[1],
                "NACIONALIDAD": valores[2],
                "TELÉFONO": valores[3],
                "DIRECCION": valores[4],
                "PLACA": valores[5],
                "DIA": fecha_dt.day,
                "MES": fecha_dt.month,
                "YEAR": fecha_dt.year,
                "DIAS2": fecha_dia_siguiente.day,
                "MES2": fecha_dia_siguiente.month,
                "YEAR2": fecha_dia_siguiente.year,
                "TARIFA": valor_cuota,
                "INICIAL": valores[11],
                "LETRAS": num2words(int(valor_cuota), lang="es").upper(),
                "TELEFONO_REF": valores[14]
            }

            ruta_pdf = generar_contrato(valores_dict)
            if ruta_pdf:
                messagebox.showinfo("Éxito", f"Cliente guardado correctamente.\nContrato generado: {ruta_pdf}")
            else:
                messagebox.showinfo("Éxito", "Cliente guardado correctamente.\n(No se generó contrato porque no se encontró la plantilla.)")

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar el cliente.\n{e}")
            return

        # Actualizar treeview
        tree.delete(*tree.get_children())
        for fila in obtener_datos_clientes():
            tree.insert("", "end", values=fila)
        ajustar_columnas(tree)

        global datos_originales
        datos_originales = [tree.item(item)["values"] for item in tree.get_children()]
        ventana_clientes.focus_force()

    def actualizar_cliente():
        # Obtener valores del formulario
        valores = {
            "Cedula": entries["Cédula"].get().strip(),
            "Nombre": entries["Nombre"].get().strip(),
            "Nacionalidad": entries["Nacionalidad"].get().strip(),
            "Telefono": entries["Teléfono"].get().strip(),
            "Direccion": entries["Dirección"].get().strip(),
            "Placa": entries["Placa"].get().strip(),
            "Fecha_inicio": convertir_fecha_formato_sqlite(entries["Fecha Inicio"].get().strip()),
            "Fecha_final": entries["Fecha Final"].get().strip(),
            "Tipo_contrato": entries["Tipo Contrato"].get().strip(),
            "Valor_cuota": entries["Valor Cuota"].get().strip(),
            "Estado": combo_estado.get().strip(),
            "Otras_deudas": entries["Otras deudas"].get().strip(),
            "Visitador": entries["Visitador"].get().strip(),
            "Referencia": entries["Referencia"].get().strip(),
            "Telefono_ref": entries["Telefono Ref"].get().strip()
        }

        # Validación de campos vacíos
        if '' in valores.values():
            messagebox.showerror("Error", "Todos los campos deben estar llenos.")
            ventana_clientes.focus_force()
            return

        try:
            with engine.begin() as conn:
                # Obtener estado y placa actuales del cliente
                stmt_check = (
                    select(
                        tabla_clientes.c.estado,
                        tabla_clientes.c.placa
                    )
                    .where(tabla_clientes.c.cedula == valores["Cedula"])
                    .limit(1)
                )
                resultado = conn.execute(stmt_check).fetchone()

                if not resultado:
                    messagebox.showerror("Error", f"No existe un cliente con la Cédula {valores['Cedula']}.")
                    return

                estado_anterior, placa_actual = resultado
                placa_actual = placa_actual.strip()

                print(f"Estado anterior: {estado_anterior}, Placa actual: {placa_actual}")

                # Lógica para cambiar placa según cambio de estado
                if estado_anterior == "activo" and valores["Estado"] == "inactivo":
                    if "**" not in placa_actual:
                        valores["Placa"] = placa_actual + " **"
                        print(f"Placa actualizada a INACTIVO: {valores['Placa']}")

                elif estado_anterior == "inactivo" and valores["Estado"] == "activo":
                    valores["Placa"] = placa_actual.replace(" **", "").strip()
                    print(f"Placa actualizada a ACTIVO: {valores['Placa']}")

                    # Validar que la placa esté en propietarios
                    stmt_prop = select(tabla_propietario.c.placa).where(tabla_propietario.c.placa == valores["Placa"]).limit(1)
                    if not conn.execute(stmt_prop).fetchone():
                        messagebox.showerror("Error", f"La Placa '{valores['Placa']}' no está registrada en propietarios.")
                        return

                    # Validar que la placa no esté ya asignada a otro cliente
                    stmt_placa = select(tabla_clientes.c.placa).where(
                        (tabla_clientes.c.placa == valores["Placa"]) &
                        (tabla_clientes.c.cedula != valores["Cedula"])
                    ).limit(1)
                    if conn.execute(stmt_placa).fetchone():
                        messagebox.showerror("Error", f"La Placa '{valores['Placa']}' ya está asignada a otro cliente.")
                        return

                # Actualizar tabla clientes
                stmt_update_cliente = (
                    update(tabla_clientes)
                    .where(tabla_clientes.c.cedula == valores["Cedula"])
                    .values(
                        nombre=valores["Nombre"],
                        nacionalidad=valores["Nacionalidad"],
                        telefono=valores["Telefono"],
                        direccion=valores["Direccion"],
                        placa=valores["Placa"],
                        fecha_inicio=valores["Fecha_inicio"],
                        fecha_final=valores["Fecha_final"],
                        tipo_contrato=valores["Tipo_contrato"],
                        valor_cuota=valores["Valor_cuota"],
                        estado=valores["Estado"],
                        otras_deudas=valores["Otras_deudas"],
                        visitador=valores["Visitador"],
                        referencia=valores["Referencia"],
                        telefono_ref=valores["Telefono_ref"]
                    )
                )
                conn.execute(stmt_update_cliente)

                # Actualizar registros relacionados
                stmt_update_registros = (
                    update(tabla_registros)
                    .where(tabla_registros.c.cedula == valores["Cedula"])
                    .values(
                        nombre=valores["Nombre"],
                        placa=valores["Placa"]
                    )
                )
                conn.execute(stmt_update_registros)

            messagebox.showinfo("Éxito", "Cliente actualizado correctamente.")
            ventana_clientes.focus_force()

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo actualizar el cliente.\n{e}")
            return

        # 🔹 Actualizar TreeView
        tree.delete(*tree.get_children())
        for fila in obtener_datos_clientes():
            tree.insert("", "end", values=fila)
        ajustar_columnas(tree)

        # 🔹 Actualizar datos_originales
        global datos_originales
        datos_originales = [tree.item(item)["values"] for item in tree.get_children()]
        ventana_clientes.focus_force()

    def cargar_datos_desde_treeview():
        # Carga los datos seleccionados del Treeview en los campos del formulario.
        seleccion = tree.selection()
        
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un cliente en la tabla.")
            return

        # Obtiene los valores de la fila seleccionada
        valores = tree.item(seleccion[0], "values")
        
        for item_id in seleccion:
            valores = tree.item(item_id, "values")
            print(valores)


        if not valores:
            messagebox.showwarning("Advertencia", "No hay datos en la selección.")
            return
        
        #print(valores)
        # Asignación de valores a cada campo del formulario
        entries["Cédula"].delete(0, tk.END)
        entries["Cédula"].insert(0, valores[0])
        entries["Nombre"].delete(0, tk.END)
        entries["Nombre"].insert(0, valores[1])
        entries["Nacionalidad"].delete(0, tk.END)
        entries["Nacionalidad"].insert(0, valores[2])
        entries["Teléfono"].delete(0, tk.END)
        entries["Teléfono"].insert(0, valores[3])
        entries["Dirección"].delete(0, tk.END)
        entries["Dirección"].insert(0, valores[4])
        entries["Placa"].delete(0, tk.END)
        entries["Placa"].insert(0, valores[5])
        entries["Fecha Inicio"].set_date(valores[8])  # Para DateEntry
        entries["Fecha Final"].insert(0, valores[9])  # Para DateEntry

        # Manejo del Combobox de "Tipo Contrato"
        opciones_tipo_contrato = entries["Tipo Contrato"]["values"]
        if valores[10] in opciones_tipo_contrato:
            entries["Tipo Contrato"].set(valores[10])
        else:
            print(f"El valor '{valores[8]}' no está en las opciones del Combobox")
        entries["Valor Cuota"].delete(0, tk.END)
        entries["Valor Cuota"].insert(0, valores[11])
        combo_estado.set(valores[12])
        entries["Otras deudas"].delete(0, tk.END)
        entries["Otras deudas"].insert(0, valores[13])
        entries["Visitador"].delete(0, tk.END)
        entries["Visitador"].insert(0, valores[14])
        entries["Referencia"].delete(0, tk.END)
        entries["Referencia"].insert(0, valores[15])
        entries["Telefono Ref"].delete(0, tk.END)
        entries["Telefono Ref"].insert(0, valores[16])

    # Configurar estilo de los botones
    style = ttk.Style()
    style.configure("TButton", font=("Arial", 12, "bold"), padding=6, width=12)
    style.configure("BotonCrear.TButton", background="#4CAF50", foreground="black")
    style.configure("BotonModificar.TButton", background="#FFC107", foreground="black")
    style.configure("BotonLimpiar.TButton", background="#F44336", foreground="black")
    # Configurar el frame con un borde
    frame_buttons = ttk.Frame(ventana_clientes, relief="ridge", borderwidth=3)
    frame_buttons.grid(row=2, column=0, columnspan=4, pady=10, padx=10, sticky="ew")
    # Botones Crear, Modificar, Limpiar con estilos personalizados
    btn_crear = ttk.Button(frame_buttons, text="Crear", command=registrar_cliente, style="BotonCrear.TButton")
    btn_crear.grid(row=0, column=0, padx=10, pady=5)
    btn_modificar = ttk.Button(frame_buttons, text="Modificar", command=actualizar_cliente, style="BotonModificar.TButton")
    btn_modificar.grid(row=0, column=1, padx=10, pady=5)
    btn_limpiar = ttk.Button(frame_buttons, text="Limpiar", command=limpiar_formulario, style="BotonLimpiar.TButton")
    btn_limpiar.grid(row=0, column=2, padx=10, pady=5)
    tree.bind("<Double-1>", lambda event: cargar_datos_desde_treeview())
    # Expansión de filas y columnas
    ventana_clientes.columnconfigure(0, weight=1)
    ventana_clientes.rowconfigure(0, weight=1)
    ventana_clientes.protocol("WM_DELETE_WINDOW", cerrar_ventana_clientes)
    return ventana_clientes  # Si quieres capturar la ventana creada

# ---------- Cerrar ventana de clientes ----------
def cerrar_ventana_clientes():
    global ventana_clientes
    ventana_clientes.destroy()
    ventana_clientes = None  # Resetea la variable

# ---------- Abrir ventana de gestión de cuentas ----------
def abrir_ventana_cuentas():
    ventana_cuentas = tk.Toplevel()
    ventana_cuentas.title("Gestión de Cuentas")
    ventana_cuentas.geometry("600x400")
    ventana_cuentas.rowconfigure(0, weight=1)
    ventana_cuentas.columnconfigure(0, weight=1)

    icono_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'img', 'inicio.ico')
    if os.path.exists(icono_path):
        ventana_cuentas.iconbitmap(icono_path)

    frame_tabla = ttk.Frame(ventana_cuentas)
    frame_tabla.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
    frame_tabla.rowconfigure(0, weight=1)
    frame_tabla.columnconfigure(0, weight=1)

    columnas = ("ID", "Nombre cuenta", "Llave")
    tree = ttk.Treeview(frame_tabla, columns=columnas, show="headings")
    scrollbar = ttk.Scrollbar(frame_tabla, orient="vertical", command=tree.yview)
    tree.configure(yscroll=scrollbar.set)

    for col in columnas:
        tree.heading(col, text=col, anchor="center")
        tree.column(col, anchor="center", width=150, stretch=True)

    tree.grid(row=0, column=0, sticky="nsew")
    scrollbar.grid(row=0, column=1, sticky="ns")

    def cargar_datos():
        try:
            tree.delete(*tree.get_children())
            with engine.connect() as conn:
                rows = conn.execute(select(
                    tabla_cuentas.c.id,
                    tabla_cuentas.c.nombre_cuenta,
                    tabla_cuentas.c.llave
                )).fetchall()

                for row in rows:
                    valores = tuple(str(c) for c in row)
                    tree.insert("", "end", values=valores)


        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron cargar los datos: {e}")
            ventana_cuentas.focus_force()

    def crear_cuenta():
        titular_valor = entry_titular.get().strip()
        llave_valor = entry_llave.get().strip()

        if not titular_valor or not llave_valor:
            messagebox.showwarning("Advertencia", "Todos los campos deben ser completados")
            ventana_cuentas.focus_force()
            return

        try:
            with engine.begin() as conn:
                # Verificar existencia
                stmt_check = select(tabla_cuentas).where(
                    (tabla_cuentas.c.nombre_cuenta == titular_valor) &
                    (tabla_cuentas.c.llave == llave_valor)
                )
                if conn.execute(stmt_check).fetchone():
                    messagebox.showwarning("Advertencia", "La combinación Titular - Llave ya existe.")
                    entry_llave.focus_force()
                    return

                # Insertar nueva cuenta
                result = conn.execute(
                    insert(tabla_cuentas).returning(tabla_cuentas.c.id),
                    {"nombre_cuenta": titular_valor, "llave": llave_valor}
                )
                new_id = result.scalar()

                tree.insert("", "end", values=(new_id, titular_valor, llave_valor))
                entry_titular.delete(0, tk.END)
                entry_llave.delete(0, tk.END)

                messagebox.showinfo("Éxito", "Cuenta creada exitosamente")
                ventana_cuentas.focus_force()

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo crear la cuenta: {e}")
            ventana_cuentas.focus_force()

    def eliminar_cuenta():
        selected_item = tree.selection()
        if not selected_item:
            messagebox.showwarning("Advertencia", "Seleccione un registro para eliminar.")
            ventana_cuentas.focus_force()
            return

        item_values = tree.item(selected_item)["values"]
        if not item_values:
            messagebox.showerror("Error", "No se pudo obtener la información del registro seleccionado.")
            return

        id_cuenta = item_values[0]
        confirmacion = messagebox.askyesno("Confirmar", f"¿Deseas eliminar la cuenta con ID {id_cuenta}?")
        ventana_cuentas.focus_force()

        if confirmacion:
            try:
                with engine.begin() as conn:
                    conn.execute(delete(tabla_cuentas).where(tabla_cuentas.c.id == id_cuenta))
                    tree.delete(selected_item)
                    messagebox.showinfo("Éxito", f"Cuenta con ID {id_cuenta} eliminada.")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo eliminar: {e}")
            ventana_cuentas.focus_force()

    cargar_datos()

    # Frame para formulario
    frame_formulario = ttk.Frame(ventana_cuentas, padding=10)
    frame_formulario.grid(row=1, column=0, columnspan=3, pady=10, sticky="ew")

    label_titular = ttk.Label(frame_formulario, text="Entidad Titular:")
    label_titular.grid(row=0, column=0, padx=5, pady=5, sticky="w")

    titular_var = tk.StringVar()
    titular_var.trace_add("write", lambda *args: titular_var.set(titular_var.get().title()))
    entry_titular = ttk.Entry(frame_formulario, textvariable=titular_var, width=30)
    entry_titular.grid(row=0, column=1, padx=5, pady=5)

    label_llave = ttk.Label(frame_formulario, text="Llave:")
    label_llave.grid(row=1, column=0, padx=5, pady=5, sticky="w")

    entry_llave = ttk.Entry(frame_formulario, width=30)
    entry_llave.grid(row=1, column=1, padx=5, pady=5)

    style = ttk.Style()
    style.configure("TButton", font=("Arial", 12, "bold"), padding=6, width=12)
    style.configure("BotonCrear.TButton", background="#4CAF50", foreground="black")
    style.configure("BotonLimpiar.TButton", background="#F44336", foreground="black")

    frame_botones = ttk.Frame(ventana_cuentas, relief="ridge", borderwidth=3)
    frame_botones.grid(row=2, column=0, columnspan=4, pady=10, padx=10, sticky="ew")

    btn_crear = ttk.Button(frame_botones, text="Crear", command=crear_cuenta, style="BotonCrear.TButton")
    btn_crear.grid(row=0, column=0, padx=10, pady=5)

    btn_eliminar = ttk.Button(frame_botones, text="Eliminar", command=eliminar_cuenta, style="BotonLimpiar.TButton")
    btn_eliminar.grid(row=0, column=1, padx=10, pady=5)

    ventana_cuentas.columnconfigure(0, weight=1)
    ventana_cuentas.rowconfigure(0, weight=1)

# ---------- Función para mostrar registros de un cliente ----------
def mostrar_registros(entry_cedula):
    cedula = entry_cedula.get().strip()
    # Obtener alto de la pantalla
    
    if not cedula:
        messagebox.showerror("Error", "Debe ingresar un cliente.")
        return
    
    def obtener_area_usable():
        user32 = ctypes.windll.user32
        rect = ctypes.wintypes.RECT()
        # SPI_GETWORKAREA = 48
        if user32.SystemParametersInfoW(48, 0, ctypes.byref(rect), 0):
            ancho = rect.right - rect.left
            alto = rect.bottom - rect.top
            return ancho, alto
        else:
            return None, None

    ventana = tk.Toplevel()
    alto_pantalla = ventana.winfo_screenheight()
    ventana.title("Extracto del Cliente")
    ancho_usable, alto_usable = obtener_area_usable()
    if ancho_usable and alto_usable:
        ventana.geometry(f"900x{alto_usable-100}")  # 800 de ancho fijo, altura usable
    else:
        ventana.geometry("800x600")  # fallback

    try:
        """Obtiene un cliente y sus pagos de la base de datos PostgreSQL."""
        
        with Session(engine) as session:
            # Obtener cliente por cédula
            cliente = session.query(
                tabla_clientes.c.cedula,
                tabla_clientes.c.nombre,
                tabla_clientes.c.placa,
                tabla_clientes.c.fecha_inicio,
                tabla_clientes.c.valor_cuota
            ).filter(tabla_clientes.c.cedula == cedula).first()

            if not cliente:
                messagebox.showerror("Error", "Cliente no encontrado.")
                return

            cedula, nombre, placa, fecha_inicio, valor_cuota = cliente

            if not fecha_inicio or not valor_cuota:
                messagebox.showerror("Error", "Datos del cliente incompletos.")
                return

            fecha_inicio = datetime.strptime(str(fecha_inicio), "%Y-%m-%d")
            fecha_actual = datetime.today()

            print(f"- Fecha inicio: {fecha_inicio} (tipo: {type(fecha_inicio)})")
            print(f"- Fecha actual: {fecha_actual} (tipo: {type(fecha_actual)})")

            # Obtener registros de pagos
            registros = session.query(
                tabla_registros.c.fecha_registro,
                tabla_registros.c.valor,
                tabla_registros.c.tipo,
                tabla_registros.c.referencia
            ).filter(tabla_registros.c.cedula == cedula).order_by(tabla_registros.c.fecha_registro).all()

            registros_modificados = []
            for fecha, valor, tipo, referencia in registros:
                fecha_dt = datetime.combine(fecha, time())
                print(f"Fecha registro: {fecha_dt} (tipo: {type(fecha_dt)})")
                registros_modificados.append((fecha_dt, valor, tipo, referencia))

            

        total = sum(row[1] for row in registros_modificados)  # row[1] es el campo "Valor"
        cuotas_pagadas = math.ceil(total / valor_cuota)  # Redondea al siguiente entero
        
        # Calcular la diferencia de días entre fecha_inicio y fecha_actual
        dias_rango = (fecha_actual - fecha_inicio).days + 1

        # Si cuotas_pagadas es mayor que la cantidad de días en el rango, ampliamos la fecha_actual
        if cuotas_pagadas > dias_rango:
            diferencia = cuotas_pagadas - dias_rango
            fecha_actual += timedelta(days=diferencia)


        # Generación del DataFrame con fechas
        fechas = pd.date_range(start=fecha_inicio, end=fecha_actual)
        df = pd.DataFrame({
            "Fecha Programada": fechas.strftime("%Y-%m-%d"),
            "Fecha Pago": "", 
            "Valor Pagado": 0, 
            "Tipo": "", 
            "Referencia": ""
        })

        # Aplicar pagos al DataFrame
        saldo = 0
        pagos_idx = 0
        for i in range(len(df)):
            while pagos_idx < len(registros_modificados) and saldo < valor_cuota:
                registro_fecha, valor, tipo, referencia = registros_modificados[pagos_idx]
                while valor + saldo >= valor_cuota:  # Manejar múltiples cuotas con un solo pago
                    falta_para_cuota = valor_cuota - saldo
                    # Evitar sobrescribir valores existentes
                    df.at[i, "Valor Pagado"] = df.at[i, "Valor Pagado"] if pd.notna(df.at[i, "Valor Pagado"]) else 0
                    df.at[i, "Valor Pagado"] += falta_para_cuota
                    df.at[i, "Fecha Pago"] = registro_fecha
                    #df.at[i, "Fecha Pago"] = registro_fecha.date().strftime("%d-%m-%Y")
                    # df.at[i, "Fecha Pago"] = datetime.strptime(registro_fecha, "%Y-%m-%d").strftime("%d-%m-%Y")
                    if pd.isna(df.at[i, "Referencia"]) or df.at[i, "Referencia"] == "":
                        df.at[i, "Referencia"] = referencia
                    if pd.isna(df.at[i, "Tipo"]) or df.at[i, "Tipo"] == "":
                        df.at[i, "Tipo"] = tipo
                    valor -= falta_para_cuota
                    saldo = 0  # Reiniciar saldo porque la cuota se completó
                    i += 1  # Pasar a la siguiente fila (día)
                    if i >= len(df):  # Evitar salir del índice
                        break

                saldo += valor
                #Solo asignar si todavía hay valor a registrar
                if valor > 0:
                    df.at[i, "Valor Pagado"] = df.at[i, "Valor Pagado"] if pd.notna(df.at[i, "Valor Pagado"]) else 0
                    df.at[i, "Valor Pagado"] += valor
                    if pd.isna(df.at[i, "Tipo"]) or df.at[i, "Tipo"] == "":
                        df.at[i, "Tipo"] = tipo
                if saldo >= valor_cuota:
                    df.at[i, "Fecha Pago"] = registro_fecha
                    #df.at[i, "Fecha Pago"] = registro_fecha.date().strftime("%d-%m-%Y")
                    # df.at[i, "Fecha Pago"] = datetime.strptime(registro_fecha, "%Y-%m-%d").strftime("%d-%m-%Y")
                    if pd.isna(df.at[i, "Referencia"]) or df.at[i, "Referencia"] == "":
                        df.at[i, "Referencia"] = referencia
                    if pd.isna(df.at[i, "Tipo"]) or df.at[i, "Tipo"] == "":
                        df.at[i, "Tipo"] = tipo
                    saldo -= valor_cuota
                else:
                    pagos_idx += 1  # Seguimos con el siguiente pago

        
        
        #df["Fecha Programada"] = pd.to_datetime(df["Fecha Programada"]).dt.strftime("%d-%m-%Y")
        df["Fecha Programada"] = pd.to_datetime(df["Fecha Programada"])
        
        # Cuotas completas pagadas
        cuotas_pagadas_completas = (df["Valor Pagado"] // valor_cuota).sum()
        # Verificar si hay un remanente en la última fila
        remanente = (df["Valor Pagado"] % valor_cuota).sum()  # Suma de los sobrantes
        # Si el remanente es suficiente para una fracción de cuota, la contamos proporcionalmente
        fraccion_cuota = remanente / valor_cuota
        # Total de cuotas pagadas (sin redondear para cálculos)
        cuotas_pagadas = cuotas_pagadas_completas + fraccion_cuota
        # Cuotas vencidas
        cuotas_vencidas = len(df)
        # Cuotas pendientes (sin redondear para cálculos)
        cuotas_pendientes = cuotas_vencidas - cuotas_pagadas
        # Valor pendiente (se mantiene con precisión completa)
        valor_pendiente = cuotas_pendientes * valor_cuota
        valor_pendiente_cop = f"${valor_pendiente:,.0f}".replace(",", ".")
        # 🔹 Solo para mostrar con 1 decimal (sin afectar cálculos internos)
        
        def capturar_y_copiar():
            # Obtener coordenadas de la ventana
            x = ventana.winfo_rootx()
            y = ventana.winfo_rooty()
            ancho = x + ventana.winfo_width()
            alto = y + ventana.winfo_height()
            # Capturar la pantalla dentro de la ventana
            captura = ImageGrab.grab(bbox=(x, y, ancho, alto))
            # Guardar la imagen en un buffer de memoria en formato PNG
            output = io.BytesIO()
            captura.save(output, format="PNG")
            image_data = output.getvalue()
            # Copiar la imagen al portapapeles en formato DIB (Device Independent Bitmap)
            output.seek(0)
            image = ImageGrab.grab(bbox=(x, y, ancho, alto)).convert("RGB")  # Convertir a RGB
            output = io.BytesIO()
            image.save(output, format="DIB")
            data = output.getvalue()
            # Abrir el portapapeles y copiar la imagen
            win32clipboard.OpenClipboard()
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
            win32clipboard.CloseClipboard()
            ventana.destroy()

        # ====== TÍTULO ======
        lbl_titulo = tk.Label(ventana, text="Extracto de pagos", font=("Arial", 16, "bold"))
        lbl_titulo.pack(pady=5)
        # ====== Botón de Acción ======
        boton_accion = tk.Button(ventana, text="Captura", font=("Arial", 10), bg="#d9d9d9", command=capturar_y_copiar)
        boton_accion.pack(pady=5)

        # ====== PANEL DE INFORMACIÓN ======
        frame_info = tk.Frame(ventana, bg="#f0f0f0", bd=2, relief="solid", padx=10, pady=10)
        frame_info.pack(fill=tk.X, padx=10, pady=5)

        # Crear los 4 paneles que representarán las columnas
        frame_col1 = tk.Frame(frame_info, bg="#f0f0f0")
        frame_col2 = tk.Frame(frame_info, bg="#f0f0f0")
        frame_col3 = tk.Frame(frame_info, bg="#f0f0f0")
        frame_col4 = tk.Frame(frame_info, bg="#f0f0f0")

        # Empaquetar los frames en línea horizontal
        frame_col1.pack(side=tk.LEFT, expand=True, padx=10, pady=5)
        frame_col2.pack(side=tk.LEFT, expand=True, padx=10, pady=5)
        frame_col3.pack(side=tk.LEFT, expand=True, padx=10, pady=5)
        frame_col4.pack(side=tk.LEFT, expand=True, padx=10, pady=5)

        # ====== Información del Cliente (Columna 1) ======
        tk.Label(frame_col1, text="Información del Cliente", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(anchor="w")
        tk.Label(frame_col1, text=f"Cédula: {cedula}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")
        tk.Label(frame_col1, text=f"Nombre: {nombre}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")

        # ====== Información del Vehículo (Columna 2) ======
        tk.Label(frame_col2, text="Información del Vehículo", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(anchor="w")
        tk.Label(frame_col2, text=f"Placa: {placa}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")
        valor_cuota_cop = f"${valor_cuota:,.0f}".replace(",", ".")
        tk.Label(frame_col2, text=f"Valor cuota: {valor_cuota_cop}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")

        # ====== Datos Financieros (Columna 3) ======
        tk.Label(frame_col3, text="Datos Financieros", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(anchor="w")
        tk.Label(frame_col3, text=f"Cuotas generadas: {cuotas_vencidas}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")
        tk.Label(frame_col3, text=f"Cuotas pagadas: {cuotas_pagadas:.1f}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")

        # ====== Estado Financiero (Columna 4) ======
        tk.Label(frame_col4, text="Estado Financiero", font=("Arial", 12, "bold"), bg="#f0f0f0").pack(anchor="w")
        tk.Label(frame_col4, text=f"Cuotas pendientes: {cuotas_pendientes:.1f}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")
        tk.Label(frame_col4, text=f"Valor para estar al día: {valor_pendiente_cop}", font=("Arial", 11), bg="#f0f0f0").pack(anchor="w")

        # ====== TREEVIEW ======
        frame_tree = tk.Frame(ventana)
        frame_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        columnas = ("Fecha Programada", "Fecha Pago", "Valor Pagado", "Tipo", "Referencia")
        # Scrollbar vertical
        scrollbar_y = ttk.Scrollbar(frame_tree, orient="vertical")
        tree = ttk.Treeview(frame_tree, columns=columnas, show='headings', style="Custom.Treeview")

        # Configurar encabezados
        for col in columnas:
            tree.heading(col, text=col, anchor="center")
            tree.column(col, anchor="center", width=150)

        # Estilos del Treeview
        style = ttk.Style()
        style.configure("Custom.Treeview", rowheight=25)
        style.configure("Treeview.Heading", font=("Arial", 10, "bold"))
        # Definir un estilo para las filas resaltadas
        tree.tag_configure("resaltado", background="lightcoral")
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        # Scrollbars
        scrollbar_y = ttk.Scrollbar(frame_tree, orient="vertical", command=tree.yview)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
        tree.configure(yscrollcommand=scrollbar_y.set)
        

        # Insertar datos en el Treeview desde el DataFrame
        for _, row in df.iterrows():
            tags = ("pendiente",) if row["Valor Pagado"] == 0 else ()
            valor_fecha_pago = row["Fecha Pago"]
            valor_formateado = ""
            if pd.notnull(valor_fecha_pago) and valor_fecha_pago:
                if isinstance(valor_fecha_pago, datetime):
                    # Si ya es un datetime, formatear
                    valor_formateado = valor_fecha_pago.strftime("%d-%m-%Y")
                else:
                    # Si es un str, parsear y formatear
                    try:
                        fecha_dt = datetime.strptime(valor_fecha_pago, "%Y-%m-%d")
                        valor_formateado = fecha_dt.strftime("%d-%m-%Y")
                    except ValueError:
                        # Si falla, deja el texto original
                        valor_formateado = valor_fecha_pago
                    
            print(f"Valor formateado: {valor_formateado} (tipo: {type(valor_formateado)})")
            # Insertar fila en el Treeview    
            tree.insert(
                "", 
                "end", 
                values=(
                    row["Fecha Programada"].strftime("%d-%m-%Y"),
                    valor_formateado,  # Aquí usamos el valor formateado
                    row["Valor Pagado"],
                    row["Tipo"],
                    row["Referencia"]
                ),
                tags=tags
            )


        # Estilo para saldo pendiente
        tree.tag_configure("pendiente", foreground="red")
        tree.yview_moveto(1.0)


    except Exception as e:
        messagebox.showerror("Error", f"Error inesperado: {str(e)}")

# ---------- Ventana de gestión de propietarios ----------
def ventana_propietario():
    # Crear ventana secundaria
    ventana_propietario = tk.Toplevel()
    ventana_propietario.title("Gestión de Propietarios")
    ventana_propietario.geometry("700x400")
    # Configurar el grid en la ventana principal
    ventana_propietario.columnconfigure(0, weight=1)
    ventana_propietario.rowconfigure(0, weight=1)
    ventana_propietario.rowconfigure(1, weight=0)
    ventana_propietario.rowconfigure(2, weight=0)
    
    def seleccionar_fila(event):
        # Obtener la fila seleccionada
        item = tree.selection()
        if item:
            valores = tree.item(item, "values")  # Obtener valores de la fila
            # Asignar los valores a las variables del formulario
            placa_var.set(valores[0])  # Placa
            modelo_var.set(valores[1])  # Modelo
            color_var.set(valores[2]) #Color
            tipo_var.set(valores[3])  # Tipo
            tarjeta_var.set(valores[4])  # Tarjeta Propiedad
            cuenta_var.set(valores[5])  # Cuenta

    def cargar_propietarios():
        """Carga los datos de la tabla 'propietario' al Treeview para filtrarlos después."""

        stmt = select(
            tabla_propietario.c.placa,
            tabla_propietario.c.modelo,
            tabla_propietario.c.color,
            tabla_propietario.c.tipo,
            tabla_propietario.c.tarjeta_propiedad,
            tabla_propietario.c.cuenta
        )

        with engine.connect() as conn:
            result = conn.execute(stmt)
            global data
            data = [tuple(row) for row in result.fetchall()]  # Asegura tuplas simples

        tree.delete(*tree.get_children())

        for fila in data:
            tree.insert("", "end", values=list(fila))  # Asegura que no sean tuplas anidadas

    def limpiar_campos():
        placa_var.set("")
        modelo_var.set("")
        color_var.set("")
        tipo_var.set("")
        tarjeta_var.set("")
        cuenta_var.set("")

    def agregar_propietario():
        """Agrega un nuevo propietario a la base de datos usando SQLAlchemy Core."""

        placa = placa_var.get().strip()
        modelo = modelo_var.get().strip()
        color = color_var.get().strip()
        tipo = tipo_var.get().strip()
        tarjeta = tarjeta_var.get().strip()
        cuenta = cuenta_var.get().strip()

        if not placa or not modelo or not tarjeta or not color or not tipo or not cuenta:
            messagebox.showerror("Error", "Todos los campos son obligatorios.")
            ventana_propietario.focus_force()
            return

        with engine.connect() as conn:
            # Verificar existencia
            stmt_verificar = select(func.count()).select_from(tabla_propietario).where(tabla_propietario.c.placa == placa)
            resultado = conn.execute(stmt_verificar).scalar()

            if resultado > 0:
                messagebox.showerror("Error", f"La placa {placa} ya está registrada.")
                ventana_propietario.focus_force()
                return

            # Insertar nuevo registro
            stmt_insertar = insert(tabla_propietario).values(
                placa=placa,
                modelo=modelo,
                color=color,
                tipo=tipo,
                tarjeta_propiedad=tarjeta,
                cuenta=cuenta
            )
            conn.execute(stmt_insertar)
            conn.commit()

        messagebox.showinfo("Éxito", "Propietario agregado correctamente.")
        cargar_propietarios()
        limpiar_campos()

    def modificar_propietario():
        """Modificar un propietario usando SQLAlchemy Core."""
        selected_item = tree.selection()
        if not selected_item:
            messagebox.showwarning("Selección requerida", "Por favor, seleccione un propietario para modificar.")
            ventana_propietario.focus_force()
            return

        placa_nueva = placa_var.get().strip().upper()
        modelo_nuevo = modelo_var.get().strip().title()
        color_nuevo = color_var.get().strip().title()
        tipo_nuevo = tipo_var.get()
        tarjeta_nueva = tarjeta_var.get().strip().title()
        cuenta_nueva = cuenta_var.get().strip().title()

        if not placa_nueva or not modelo_nuevo or not color_nuevo or not tipo_nuevo or not tarjeta_nueva or not cuenta_nueva:
            messagebox.showwarning("Campos vacíos", "Todos los campos deben estar llenos.")
            ventana_propietario.focus_force()
            return

        item = tree.item(selected_item)
        placa_original = item["values"][0]

        with engine.begin() as conn:  # Usa transacción automática
            try:
                # Verificar duplicado (otra placa igual)
                stmt_check = select(func.count()).select_from(tabla_propietario).where(
                    (tabla_propietario.c.placa == placa_nueva) & (tabla_propietario.c.placa != placa_original)
                )
                duplicados = conn.execute(stmt_check).scalar()
                if duplicados > 0:
                    messagebox.showerror("Error de duplicado", f"La placa '{placa_nueva}' ya existe en otro registro.")
                    ventana_propietario.focus_force()
                    return

                # Actualizar en tabla propietario
                stmt_update_prop = (
                    update(tabla_propietario)
                    .where(tabla_propietario.c.placa == placa_original)
                    .values(
                        placa=placa_nueva,
                        modelo=modelo_nuevo,
                        color=color_nuevo,
                        tipo=tipo_nuevo,
                        tarjeta_propiedad=tarjeta_nueva,
                        cuenta=cuenta_nueva
                    )
                )
                conn.execute(stmt_update_prop)

                # Actualizar placa en las tablas relacionadas
                conn.execute(update(tabla_clientes).where(tabla_clientes.c.placa == placa_original).values(placa=placa_nueva))
                conn.execute(update(tabla_registros).where(tabla_registros.c.placa == placa_original).values(placa=placa_nueva))

                messagebox.showinfo("Éxito", "El propietario ha sido modificado correctamente.")
                ventana_propietario.focus_force()

                # Refrescar Treeview localmente (opcional)
                tree.item(selected_item, values=(placa_nueva, modelo_nuevo, color_nuevo, tipo_nuevo, tarjeta_nueva, cuenta_nueva))
                tree.selection_set(selected_item)

            except Exception as e:
                messagebox.showerror("Error de base de datos", f"Ocurrió un error: {e}")
                ventana_propietario.focus_force()
                return

        # Recargar y limpiar
        cargar_propietarios()
        limpiar_campos()


    # 🌟 FRAME PARA EL TREEVIEW
    frame_tree = ttk.Frame(ventana_propietario, padding=10)
    frame_tree.grid(row=0, column=0, sticky="nsew")
    columnas = ("Placa", "Modelo", "Color", "Tipo", "Tarjeta Propiedad", "Cuenta")
    tree = ttk.Treeview(frame_tree, columns=columnas, show="headings", height=8)
    # Variable para almacenar el estado de orden (ascendente/descendente)
    sort_states = {col: False for col in columnas}
    
    def ordenar_por_columna(tree, col):
        """Ordena el Treeview al hacer clic en un encabezado."""
        # Obtener datos actuales del Treeview
        datos = [(tree.set(item, col), item) for item in tree.get_children('')]
        
        # Determinar el tipo de datos y ordenar correctamente
        try:
            datos.sort(key=lambda x: int(x[0]), reverse=sort_states[col])
        except ValueError:
            datos.sort(key=lambda x: x[0], reverse=sort_states[col])

        # Invertir el estado de orden para el próximo clic
        sort_states[col] = not sort_states[col]

        # Reorganizar los datos en el Treeview
        for index, (_, item) in enumerate(datos):
            tree.move(item, '', index)

        # Actualizar visualmente el encabezado (opcional)
        tree.heading(col, text=f"{col} {'▲' if sort_states[col] else '▼'}")

    # Encabezados del Treeview con eventos de clic
    for col in columnas:
        tree.heading(col, text=col, anchor="center", command=lambda c=col: ordenar_por_columna(tree, c))
        tree.column(col, width=180, anchor="center")
    # Barra de desplazamiento vertical
    scrollbar_vertical = ttk.Scrollbar(frame_tree, orient="vertical", command=tree.yview)
    scrollbar_vertical.grid(row=0, column=1, sticky="ns")
    # Configuración para que el Treeview use el scroll
    tree.configure(yscrollcommand=scrollbar_vertical.set)
    tree.grid(row=0, column=0, sticky="nsew")
    tree.bind("<Double-1>", seleccionar_fila)
    # Expandir el TreeView dentro del frame
    frame_tree.columnconfigure(0, weight=1)
    frame_tree.rowconfigure(0, weight=1)

    # 🌟 FRAME PARA FORMULARIO
    frame_form = ttk.Frame(ventana_propietario, padding=10, borderwidth=2, relief="solid")
    frame_form.grid(row=1, column=0, sticky="ew")
    
    def update_total_entries():
        """
        Cuenta la cantidad de registros actuales en el Treeview y actualiza el valor en el Entry correspondiente.
        """
        total = len(tree.get_children())
        entry_total_placas.delete(0, "end")  # Limpiar el Entry
        entry_total_placas.insert(0, str(total))  # Insertar el total
    
    def filter_treeview(*args):
        """Filtra el Treeview en función del texto ingresado en el Entry."""
        search_term = placa_var.get().lower()
        
        # Restaurar datos si no hay texto
        tree.delete(*tree.get_children())  # Limpiar Treeview
        
        for item in data:
            if any(search_term in str(value).lower() for value in item):  
                tree.insert("", "end", values=item)  # Insertar si coincide
        update_total_entries()  # Actualizar el total de registros
                
    def filter2_treeview(*args):
        """Filtra el Treeview en función del texto ingresado en el Entry."""
        search_term = tarjeta_var.get().title()
        
        # Restaurar datos si no hay texto
        tree.delete(*tree.get_children())  # Limpiar Treeview
        
        for item in data:
            if any(search_term in str(value).title() for value in item):  
                tree.insert("", "end", values=item)  # Insertar si coincide  
        update_total_entries()  # Actualizar el total de registros          
                

    # Crear campos del formulario en línea horizontal
    placa_var = tk.StringVar()
    placa_var.trace_add("write", filter_treeview)  # Ejecuta la función en cada cambio
    placa_var.trace_add("write", lambda *args: placa_var.set(placa_var.get().upper()))
    ttk.Label(frame_form, text="Placa:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
    entry_placa = ttk.Entry(frame_form, textvariable=placa_var, width=30)
    entry_placa.grid(row=0, column=1, padx=5, pady=5)

    modelo_var = tk.StringVar()
    modelo_var.trace_add("write", lambda *args: modelo_var.set(modelo_var.get().title()))
    ttk.Label(frame_form, text="Modelo:").grid(row=0, column=2, padx=5, pady=5, sticky="w")
    entry_modelo = ttk.Entry(frame_form, textvariable=modelo_var, width=30)
    entry_modelo.grid(row=0, column=3, padx=5, pady=5)
    
    color_var = tk.StringVar()
    color_var.trace_add("write", lambda *args: color_var.set(color_var.get().title()))
    ttk.Label(frame_form, text="Color:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
    entry_color = ttk.Entry(frame_form, textvariable=color_var, width=30)
    entry_color.grid(row=1, column=1, padx=5, pady=5)
    
    tipo_var = tk.StringVar()
    # Definición del Combobox con las dos opciones
    ttk.Label(frame_form, text="Tipo:").grid(row=1, column=2, padx=5, pady=5, sticky="w")
    combo_tipo = ttk.Combobox(frame_form, textvariable=tipo_var, values=["Nueva", "Usada"], state="readonly", width=28)
    combo_tipo.grid(row=1, column=3, padx=5, pady=5)
    # Seleccionar un valor por defecto (opcional)
    combo_tipo.current(0)  # Selecciona "Nueva" por defecto

    tarjeta_var = tk.StringVar()
    tarjeta_var.trace_add("write", filter2_treeview)  # Ejecuta la función en cada cambio
    tarjeta_var.trace_add("write", lambda *args: tarjeta_var.set(tarjeta_var.get().title()))
    ttk.Label(frame_form, text="Tarjeta Propiedad:").grid(row=1, column=4, padx=5, pady=5, sticky="w")
    entry_tarjeta = ttk.Entry(frame_form, textvariable=tarjeta_var, width=30)
    entry_tarjeta.grid(row=1, column=5, padx=5, pady=5)
    cuenta_var = tk.StringVar()
    cuenta_var.trace_add("write", lambda *args: cuenta_var.set(cuenta_var.get().title()))
    ttk.Label(frame_form, text="Cuenta:").grid(row=0, column=4, padx=5, pady=5, sticky="w")
    entry_cuenta = ttk.Entry(frame_form, textvariable=cuenta_var, width=30)
    entry_cuenta.grid(row=0, column=5, padx=5, pady=5)
    
    frame_info = ttk.Frame(ventana_propietario, padding=10)
    frame_info.grid(row=2, column=0, sticky="nsew")
    # Campo 1: Total Placas
    label_total_placas = ttk.Label(frame_info, text="Total Placas:")
    label_total_placas.grid(row=0, column=0, padx=5, pady=5, sticky="w")
    entry_total_placas = ttk.Entry(frame_info)
    entry_total_placas.grid(row=0, column=1, padx=5, pady=5)

    # 🌟 FRAME PARA BOTONES
    frame_buttons = ttk.Frame(ventana_propietario, relief="ridge", borderwidth=3)
    frame_buttons.grid(row=3, column=0, columnspan=4, pady=10, padx=10, sticky="ew")

    style = ttk.Style()
    style.configure("TButton", font=("Arial", 12, "bold"), padding=6, width=12)
    style.configure("BotonCrear.TButton", background="#4CAF50", foreground="black")
    style.configure("BotonModificar.TButton", background="#FFC107", foreground="black")
    style.configure("BotonLimpiar.TButton", background="#F44336", foreground="black")
    style.configure("BotonDashBoard.TButton", background="#F44336", foreground="black")
    
    # Botones con funcionalidades
    btn_crear = ttk.Button(frame_buttons, text="Crear", command= agregar_propietario, style="BotonCrear.TButton")
    btn_crear.grid(row=0, column=0, padx=10, pady=5)
    btn_modificar = ttk.Button(frame_buttons, text="Modificar", command= modificar_propietario, style="BotonModificar.TButton")
    btn_modificar.grid(row=0, column=1, padx=10, pady=5)
    btn_limpiar = ttk.Button(frame_buttons, text="Limpiar", command= limpiar_campos, style="BotonLimpiar.TButton")
    btn_limpiar.grid(row=0, column=2, padx=10, pady=5)
    
    # Expandir columnas en el frame de botones
    frame_buttons.columnconfigure((0, 1, 2), weight=1)

    cargar_propietarios()

# ---------- Función para obtener datos cuadre del dia----------
def obtener_datos(fecha_inicio, fecha_fin):
    try:
        stmt = (
            select(
                tabla_registros.c.nombre_cuenta.label("nombre_cuenta"),
                case(
                    (tabla_registros.c.motivo == 'N-a', 'Tarifas'),
                    else_=tabla_registros.c.motivo
                ).label("motivo"),
                tabla_registros.c.valor.label("valor"),
                tabla_registros.c.saldos.label("saldos")
            )
            .where(
                and_(
                    tabla_registros.c.fecha_registro >= fecha_inicio,
                    tabla_registros.c.fecha_registro <= fecha_fin
                )
            )
        )
        with engine.connect() as conn:
            df = pd.read_sql(stmt, conn)
        return df

    except Exception as e:
        print(f"💥 Error al obtener datos: {e}")
        return pd.DataFrame()

# ---------- Lógica de resumen integrada ----------
def generar_resumen_por_cuenta(df):
    resultado = []
    total_general_valor = 0
    total_general_saldos = 0

    cuentas = df['nombre_cuenta'].unique()

    for cuenta in cuentas:
        df_cuenta = df[df['nombre_cuenta'] == cuenta]

        # ---------- Fila de "Tarifas" ----------
        total_valor = df_cuenta['valor'].sum()
        resultado.append((
            cuenta,
            "Tarifas",
            f"{total_valor:,.0f}",
            "0",
            f"{total_valor:,.0f}"
        ))

        # ---------- Motivos desde columna 'motivo' pero sumando solo saldos ----------
        df_saldos_motivos = (
            df_cuenta.groupby("motivo")['saldos']
            .sum()
            .reset_index()
        )

        for _, row in df_saldos_motivos.iterrows():
            if row["motivo"] != "Tarifas":  # Evitar duplicar la fila de Tarifas
                resultado.append((
                    cuenta,
                    row["motivo"],
                    "0",
                    f"{row['saldos']:,.0f}",
                    f"{row['saldos']:,.0f}"
                ))

        # ---------- Subtotales de la cuenta ----------
        subtotal_valor = total_valor
        subtotal_saldos = df_saldos_motivos['saldos'].sum()
        subtotal_total = subtotal_valor + subtotal_saldos

        resultado.append((
            cuenta,
            "TOTAL CUENTA",
            f"{subtotal_valor:,.0f}",
            f"{subtotal_saldos:,.0f}",
            f"{subtotal_total:,.0f}"
        ))
        resultado.append(("", "", "", "", ""))  # Línea vacía

        total_general_valor += subtotal_valor
        total_general_saldos += subtotal_saldos

    # ---------- Total general ----------
    total_general = total_general_valor + total_general_saldos
    resultado.append((
        "TOTAL GENERAL",
        "",
        f"{total_general_valor:,.0f}",
        f"{total_general_saldos:,.0f}",
        f"{total_general:,.0f}"
    ))
    return resultado

# ---------- Crear interfaz ----------------------
def crear_resumen_por_cuenta_y_motivo():
    ventana = tk.Tk()
    ventana.title("Resumen por Cuenta y Motivo")
    ventana.geometry("1000x650")

    lbl_titulo = tk.Label(ventana, text="", font=("Arial", 16, "bold"))
    lbl_titulo.pack(pady=10)

    frame_top = tk.Frame(ventana)
    frame_top.pack()

    tk.Label(frame_top, text="Desde:", font=("Arial", 12)).pack(side="left", padx=5)
    fecha_inicio = DateEntry(frame_top, width=12, background='darkblue', foreground='white',
                             borderwidth=2, date_pattern='yyyy-mm-dd')
    fecha_inicio.set_date(datetime.now())
    fecha_inicio.pack(side="left")

    tk.Label(frame_top, text="Hasta:", font=("Arial", 12)).pack(side="left", padx=5)
    fecha_fin = DateEntry(frame_top, width=12, background='darkblue', foreground='white',
                          borderwidth=2, date_pattern='yyyy-mm-dd')
    fecha_fin.set_date(datetime.now())
    fecha_fin.pack(side="left")

    btn_cargar = tk.Button(frame_top, text="Cargar Resumen")
    btn_cargar.pack(side="left", padx=10, pady=5)

    btn_captura = tk.Button(frame_top, text="📸 Capturar")
    btn_captura.pack(side="left", padx=10, pady=5)

    tree = ttk.Treeview(ventana, columns=["Cuenta", "Motivo", "Total Valor", "Total Saldos", "TOTAL"], show="headings")
    for col in ["Cuenta", "Motivo", "Total Valor", "Total Saldos", "TOTAL"]:
        tree.heading(col, text=col)
        tree.column(col, anchor="center", width=180)
    tree.pack(fill="both", expand=True)

    scrollbar_y = ttk.Scrollbar(ventana, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=scrollbar_y.set)
    scrollbar_y.pack(side="right", fill="y")

    style = ttk.Style()
    style.configure("Treeview.Heading", font=("Arial", 11, "bold"))
    style.configure("Treeview", font=("Arial", 10), rowheight=25)
    tree.tag_configure("bold", font=("Arial", 10, "bold"))
    tree.tag_configure("total_general", background="#d1ffd1", font=("Arial", 11, "bold"))

    def cargar_datos():
        tree.delete(*tree.get_children())
        inicio = fecha_inicio.get_date()
        fin = fecha_fin.get_date()

        if inicio > fin:
            messagebox.showwarning("Fechas inválidas", "La fecha de inicio no puede ser posterior a la fecha final.")
            return

        lbl_titulo.config(text=f"📋 Reporte de valores del {inicio.strftime('%d-%m-%Y')} al {fin.strftime('%d-%m-%Y')}")

        df = obtener_datos(inicio, fin)
        if df.empty:
            tree.insert("", "end", values=("Sin datos", "", "", ""))
            return

        resumen = generar_resumen_por_cuenta(df)

        for fila in resumen:
            if "TOTAL CUENTA" in fila[1]:
                tree.insert("", "end", values=fila, tags=("bold",))
            elif "TOTAL GENERAL" in fila[0]:
                tree.insert("", "end", values=fila, tags=("total_general",))
            else:
                tree.insert("", "end", values=fila)

    btn_cargar.config(command=cargar_datos)

    def capturar_ventana():
        ventana.update()
        x = ventana.winfo_rootx()
        y = ventana.winfo_rooty()
        w = ventana.winfo_width()
        h = ventana.winfo_height()
        imagen = ImageGrab.grab(bbox=(x, y, x + w, y + h)).convert("RGB")
        output = io.BytesIO()
        imagen.save(output, format="BMP")
        data = output.getvalue()[14:]
        win32clipboard.OpenClipboard()
        win32clipboard.EmptyClipboard()
        win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
        win32clipboard.CloseClipboard()
        messagebox.showinfo("Captura", "📸 Captura copiada al portapapeles.")

    btn_captura.config(command=capturar_ventana)

    ventana.mainloop()

# ---------- Función para generar reporte de atrasos ----------
def reporte_atrasos():
    try:
        with engine.connect() as conn:
            df_registros = pd.read_sql(select(tabla_registros), conn)
            df_clientes = pd.read_sql(select(tabla_clientes), conn)

        df_registros.columns = df_registros.columns.str.lower()
        df_clientes.columns = df_clientes.columns.str.lower()

        hoy = datetime.now().date()
        resultados = []

        for _, cliente in df_clientes.iterrows():
            placa = cliente.get('placa', '')
            if '*' in str(placa):
                continue

            try:
                cedula = cliente['cedula']
                nombre = cliente['nombre']
                fecha_inicio = pd.to_datetime(cliente['fecha_inicio']).date()
                fecha_final = int(cliente['fecha_final'])
                valor_cuota = float(cliente['valor_cuota'])

                if valor_cuota <= 0:
                    continue

                dias_transcurridos = min((hoy - fecha_inicio).days + 1, fecha_final)
                monto_esperado = dias_transcurridos * valor_cuota

                pagos_cliente = df_registros[df_registros['cedula'] == cedula].copy()
                pagos_cliente['fecha_sistema'] = pd.to_datetime(pagos_cliente['fecha_sistema'], errors='coerce').dt.date
                pagos_cliente.dropna(subset=['fecha_sistema'], inplace=True)

                total_pagado = pagos_cliente['valor'].sum()
                dias_cubiertos = round(total_pagado / valor_cuota, 1)
                dias_atraso = dias_transcurridos - dias_cubiertos

                pagos_dias = [
                    int(pagos_cliente[pagos_cliente['fecha_sistema'] == hoy - timedelta(days=i)]['valor'].sum() / 1000)
                    for i in range(10)
                ]

                resultados.append({
                    "Cedula": cedula,
                    "Placa": placa,
                    "Nombre": nombre,
                    "Antigüedad": dias_transcurridos,
                    "Días de Atraso": round(dias_atraso, 1),
                    "Monto Adeudado": int(round(monto_esperado - total_pagado)),
                    **{f"Día {i+1}": valor for i, valor in enumerate(pagos_dias)}
                })

            except (ValueError, TypeError, KeyError):
                continue

        df = pd.DataFrame(resultados)

        if not df.empty:
            df.sort_values(by="Días de Atraso", ascending=False, inplace=True)
            df.insert(0, "#", range(1, len(df) + 1))

            # Agregar fila vacía
            df.loc[len(df)] = [""] * len(df.columns)

            # Agregar fila TOTAL
            total_adeudado = pd.to_numeric(df["Monto Adeudado"], errors='coerce').fillna(0).astype(int).sum()
            fila_total = {col: "" for col in df.columns}
            fila_total["Nombre"] = "TOTAL"
            fila_total["Monto Adeudado"] = total_adeudado
            fila_total["#"] = ""
            df.loc[len(df)] = fila_total

        return df

    except Exception as e:
        print(f"💥 Error en reporte_atrasos: {e}")
        return pd.DataFrame()

# ---------- Crear interfaz de atrasos ----------
def crear_interfaz_atrasos(root_padre, entry_cedula, entry_nombre, entry_placa):
    global ventana_atrasos

    if ventana_atrasos and ventana_atrasos.winfo_exists():
        ventana_atrasos.lift()
        return

    df = reporte_atrasos()
    if df.empty:
        print("No hay datos para mostrar.")
        return

    def formatear_monto(x):
        try:
            return locale.currency(float(x), grouping=True)
        except:
            return x

    df["Monto Adeudado"] = df["Monto Adeudado"].apply(formatear_monto)

    columnas = list(df.columns)

    def copiar_placa_al_portapapeles(event):
        selected_item = tree.focus()
        if not selected_item:
            return
        valores = tree.item(selected_item, 'values')
        if not valores:
            return
        placa = valores[columnas.index("Placa")]
        nombre = valores[columnas.index("Nombre")]
        cedula = valores[columnas.index("Cedula")]

        entry_cedula.delete(0, tk.END)
        entry_cedula.insert(0, cedula)
        entry_nombre.delete(0, tk.END)
        entry_nombre.insert(0, nombre)
        entry_placa.delete(0, tk.END)
        entry_placa.insert(0, placa)

        if len(placa) == 6:
            placa_modificada = placa[:3] + "-" + placa[3:]
            root_padre.clipboard_clear()
            root_padre.clipboard_append(placa_modificada)
            root_padre.update()
            print(f"Placa copiada al portapapeles: {placa_modificada}")

    def copiar_mensaje_personalizado(event):
        selected_item = tree.focus()
        if not selected_item:
            return
        valores = tree.item(selected_item, 'values')
        if not valores:
            return
        try:
            nombre_completo = valores[columnas.index("Nombre")]
            primer_nombre = nombre_completo.split()[0]
            antiguedad = int(valores[columnas.index("Antigüedad")])
            atraso = float(valores[columnas.index("Días de Atraso")])
            monto = valores[columnas.index("Monto Adeudado")]
            gabela = (antiguedad / 30) * 1.5

            if antiguedad < 30:
                if atraso >= 5:
                    mensaje = (
                        f"{primer_nombre}, actualmente presenta {atraso:.1f} días de atraso. "
                        "Dado que el contrato es reciente, se ha programado una visita de seguimiento por parte del equipo de cobros."
                    )
                else:
                    mensaje = (
                        f"{primer_nombre}, lleva {atraso:.1f} días de atraso. "
                        "Por favor, recuerde mantenerse al día con su plan de pago."
                    )
            elif atraso > gabela + 5:
                mensaje = (
                    f"{primer_nombre}, registra {atraso:.1f} días de atraso en los pagos. "
                    f"Este valor supera el límite permitido según la antigüedad de su contrato ({gabela:.1f} días). "
                    "Se ha programado una visita por parte del personal de cobradores."
                )
            elif atraso > gabela + 2:
                mensaje = (
                    f"{primer_nombre}, lleva {atraso:.1f} días de atraso. "
                    f"Supera el límite tolerado de {gabela:.1f} días según su antigüedad. "
                    "Lo invitamos a ponerse al día lo antes posible para evitar medidas adicionales."
                )
            else:
                mensaje = (
                    f"{primer_nombre}, tiene {atraso:.1f} días de atraso. "
                    "Le recordamos la importancia de mantener los pagos al día. "
                    f"El saldo actual pendiente es de {monto}."
                )

            root_padre.clipboard_clear()
            root_padre.clipboard_append(mensaje)
            root_padre.update()
            print("Mensaje copiado al portapapeles.")
        except Exception as e:
            print(f"Error al generar mensaje: {e}")

    ventana_atrasos = tk.Toplevel(root_padre)
    ventana_atrasos.title("Reporte de Atrasos")
    ventana_atrasos.geometry("1200x600")

    frame_principal = tk.Frame(ventana_atrasos)
    frame_principal.grid(row=0, column=0, sticky="nsew")
    ventana_atrasos.grid_rowconfigure(0, weight=1)
    ventana_atrasos.grid_columnconfigure(0, weight=1)

    entry_filtro = tk.Entry(frame_principal, font=("Arial", 12))
    entry_filtro.grid(row=0, column=0, columnspan=2, sticky="ew", padx=10, pady=5)

    tree = ttk.Treeview(frame_principal, columns=columnas, show='headings')
    style = ttk.Style()
    style.configure("Treeview", font=("Arial", 10))
    style.configure("Treeview.Heading", font=("Arial", 10, "bold"))
    tree.tag_configure('total', font=('Arial', 10, 'bold'))
    tree.tag_configure('grave', background='pink')
    tree.bind("<Double-1>", copiar_placa_al_portapapeles)
    tree.bind("<Button-3>", copiar_mensaje_personalizado)

    for col in columnas:
        tree.heading(col, text=col, anchor='center')
        tree.column(col, anchor='center', width=120)

    tree.grid(row=1, column=0, columnspan=2, sticky="nsew")
    frame_principal.grid_rowconfigure(1, weight=1)
    frame_principal.grid_columnconfigure(0, weight=1)

    scrollbar_y = ttk.Scrollbar(frame_principal, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=scrollbar_y.set)
    scrollbar_y.grid(row=1, column=2, sticky="ns")

    df_original = df.copy()

    def cargar_tree(df_view):
        tree.delete(*tree.get_children())
        for _, fila in df_view.iterrows():
            valores = list(fila)
            if str(fila["Nombre"]).strip().upper() == "TOTAL":
                tags = ('total',)
            elif isinstance(fila["Días de Atraso"], (int, float)) and fila["Días de Atraso"] > 10:
                tags = ('grave',)
            else:
                tags = ()
            tree.insert("", "end", values=valores, tags=tags)

    def aplicar_filtro(event=None):
        filtro = entry_filtro.get().lower().strip()
        if not filtro:
            cargar_tree(df_original)
            return
        filtrado = df_original[
            df_original["Nombre"].str.lower().str.contains(filtro) |
            df_original["Placa"].str.lower().str.contains(filtro)
        ]
        cargar_tree(filtrado)

    cargar_tree(df_original)
    entry_filtro.bind("<KeyRelease>", aplicar_filtro)

    def exportar_excel():
        ruta = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if ruta:
            df_original.to_excel(ruta, index=False)
            print(f"Exportado a {ruta}")

    def generar_nuevo_tree():
        seleccion = tree.selection()
        if not seleccion:
            return
        nueva_ventana = tk.Toplevel(ventana_atrasos)
        nueva_ventana.title("Selección filtrada")
        nuevo_tree = ttk.Treeview(nueva_ventana, columns=columnas, show='headings')
        for col in columnas:
            nuevo_tree.heading(col, text=col)
            nuevo_tree.column(col, anchor="center", width=120)
        nuevo_tree.pack(fill="both", expand=True)
        for item in seleccion:
            valores = tree.item(item, "values")
            nuevo_tree.insert("", "end", values=valores)

    btn_exportar = tk.Button(frame_principal, text="Exportar a Excel", command=exportar_excel)
    btn_exportar.grid(row=2, column=0, pady=5, sticky="e", padx=10)

    btn_generar = tk.Button(frame_principal, text="Reporte recogidas", command=generar_nuevo_tree)
    btn_generar.grid(row=2, column=1, pady=5, sticky="w", padx=10)


# ---------- Función para ordenar columnas en TreeView ----------
def ordenar_por_columna(tree, col, descendente):
    datos = [(tree.set(k, col), k) for k in tree.get_children('')]

    try:
        datos.sort(key=lambda t: float(t[0]), reverse=descendente)
    except ValueError:
        datos.sort(key=lambda t: t[0], reverse=descendente)

    for index, (val, k) in enumerate(datos):
        tree.move(k, '', index)

    # Cambiar orden para el próximo clic
    tree.heading(col, command=lambda: ordenar_por_columna(tree, col, not descendente))

def cargar_nombres_columnas():
    if not os.path.exists(JSON_PATH):
        columnas = ["Fecha"] + [f"Columna {i}" for i in range(1, 9)] + ["Total"]
        guardar_nombres_columnas(columnas)
    else:
        with open(JSON_PATH, 'r') as file:
            data = json.load(file)
            columnas = data.get("columnas", ["Fecha"] + [f"Columna {i}" for i in range(1, 9)] + ["Total"])
    return columnas

def guardar_nombres_columnas(columnas):
    with open(JSON_PATH, 'w') as file:
        json.dump({"columnas": columnas}, file, indent=4)

def generar_fechas():
    fecha_inicio = datetime(2025, 5, 1)
    fecha_actual = datetime.now().date()
    delta = fecha_actual - fecha_inicio
    fechas = [(fecha_inicio + timedelta(days=i)).strftime("%d-%m-%Y") for i in range(delta.days + 1)]
    return fechas

def inicializar_excel():
    if not os.path.exists(XLSX_PATH):
        wb = Workbook()
        ws = wb.active
        ws.title = "Datos"
        columnas = cargar_nombres_columnas()
        ws.append(columnas)
        wb.save(XLSX_PATH)

def cargar_datos_desde_excel(tree):
    if os.path.exists(XLSX_PATH):
        df = pd.read_excel(XLSX_PATH, sheet_name="Datos")
        if not df.empty:
            for index, row in df.iterrows():
                valores = row.tolist()
                tree.insert("", "end", values=valores, tags=('par' if index % 2 == 0 else 'impar'))
        else:
            fechas = generar_fechas()
            for i, fecha in enumerate(fechas):
                valores = [fecha] + ["0"] * 8 + ["$0"]
                tree.insert("", "end", values=valores, tags=('par' if i % 2 == 0 else 'impar'))

def guardar_en_excel(tree):
    columnas = cargar_nombres_columnas()
    data = [columnas]

    for item in tree.get_children():
        valores = tree.item(item, 'values')
        data.append(valores)

    df = pd.DataFrame(data[1:], columns=data[0])
    df.to_excel(XLSX_PATH, sheet_name="Datos", index=False)

class EditableTreeview(ttk.Treeview):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.columnas = cargar_nombres_columnas()
        self["columns"] = self.columnas
        self["show"] = "headings"

        # Estilo de encabezados
        style = ttk.Style()
        style.configure("Treeview.Heading", background="#4A90E2", foreground="black", font=("Helvetica", 10, "bold"))
        style.configure("Treeview", rowheight=25)

        # Estilo de filas alternas
        self.tag_configure('par', background="#F0F0F0")
        self.tag_configure('impar', background="#FFFFFF")
        self.tag_configure('negrita', font=('Helvetica', 10, 'bold'))

        for col in self.columnas:
            self.heading(col, text=col)
            ancho = 100 if col != "Fecha" else 120
            self.column(col, width=ancho, anchor='center')
        
        # Estilo para la última columna (Total)
        self.column(self.columnas[-1], anchor='e')
        
        vsb = ttk.Scrollbar(parent, orient="vertical", command=self.yview)
        hsb = ttk.Scrollbar(parent, orient="horizontal", command=self.xview)
        self.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")
        self.pack(expand=True, fill='both')
        self.bind("<Double-1>", self.start_edit)

    def start_edit(self, event):
        item = self.identify_row(event.y)
        column = self.identify_column(event.x)

        if not item or column == '#1' or column == f"#{len(self.columnas)}":
            return

        col_index = int(column.replace("#", "")) - 1
        x, y, width, height = self.bbox(item, column)
        entry = tk.Entry(self.parent)
        entry.place(x=x, y=y + self.winfo_rooty() - self.parent.winfo_rooty(), width=width, height=height)
        entry.insert(0, self.item(item, 'values')[col_index])
        entry.focus()

        def on_focus_out(_):
            nuevo_valor = entry.get()
            if nuevo_valor.replace('.', '', 1).isdigit():
                valores = list(self.item(item, 'values'))
                valores[col_index] = nuevo_valor
                self.item(item, values=valores)
                self.actualizar_total(item)
            entry.destroy()

        entry.bind("<FocusOut>", on_focus_out)

    def actualizar_total(self, item):
        valores = list(self.item(item, 'values'))
        try:
            suma = sum(float(valores[i]) for i in range(1, 9))
            valores[-1] = f"${int(suma):,}"
            self.item(item, values=valores, tags=('negrita',))
        except ValueError:
            print("❌ Error al intentar sumar los valores.")

def iniciar_interfaz():
    # Configuración de la ventana principal
    root = tk.Tk()
    root.title("TreeView Editable - Estilo Excel")
    root.geometry("1200x600")

    tree = EditableTreeview(root)

    btn_guardar = tk.Button(root, text="Guardar en Excel", command=lambda: guardar_en_excel(tree))
    btn_guardar.pack(pady=10)

    inicializar_excel()
    cargar_datos_desde_excel(tree)

    root.mainloop()


# ---------- Función para iniciar la ventana de deudas ----------
def iniciar_ventana_deudas():
    
        # --- FUNCIONES ---
    def cargar_clientes(filtro_nombre="", filtro_placa=""):
        tree_clientes.delete(*tree_clientes.get_children())
        with engine.begin() as conn:
            stmt = select(tabla_clientes.c.nombre, tabla_clientes.c.placa, tabla_clientes.c.cedula)

            condiciones = []
            if filtro_nombre:
                condiciones.append(tabla_clientes.c.nombre.ilike(f"%{filtro_nombre}%"))
            if filtro_placa:
                condiciones.append(tabla_clientes.c.placa.ilike(f"%{filtro_placa}%"))

            if condiciones:
                stmt = stmt.where(and_(*condiciones))

            stmt = stmt.order_by(tabla_clientes.c.placa)
            for nombre, placa, cedula in conn.execute(stmt):
                tree_clientes.insert("", "end", values=(nombre, placa), tags=(cedula, placa, nombre))

    def actualizar_entry(entry, valor):
        entry.config(state="normal")
        entry.delete(0, tk.END)
        entry.insert(0, f"{valor:.0f}")
        entry.config(state="readonly")

    def cargar_deudas(cedula, placa, entry_deudas, entry_actual, entry_abonos):
        total = 0
        tree_deudas.delete(*tree_deudas.get_children())
        with engine.begin() as conn:
            stmt = select(tabla_otras_deudas).where(
                tabla_otras_deudas.c.cedula == cedula,
                tabla_otras_deudas.c.placa == placa
            ).order_by(tabla_otras_deudas.c.fecha_deuda.desc())
            for row in conn.execute(stmt):
                tree_deudas.insert("", "end", values=(
                    row.id, row.fecha_deuda, row.descripcion, row.valor
                ))
                total += float(row.valor)
        actualizar_entry(entry_deudas, total)
        print(f"Total Deudas: {total:.0f}")
        actualizar_resumen(entry_deudas, entry_abonos, entry_actual)

    def cargar_abonos(nombre, placa, entry_abonos, entry_actual, entry_deudas):
        total = 0
        tree_abonos.delete(*tree_abonos.get_children())
        with engine.begin() as conn:
            stmt = select(
                tabla_registros.c.fecha_registro,
                tabla_registros.c.saldos,
                tabla_registros.c.motivo,
                tabla_registros.c.tipo
            ).where(
                and_(
                    tabla_registros.c.nombre == nombre,
                    tabla_registros.c.placa == placa,
                    tabla_registros.c.saldos > 0,
                    tabla_registros.c.motivo.notin_(["N-a", "Multa"])
                )
            ).order_by(tabla_registros.c.fecha_registro.desc())
            for row in conn.execute(stmt):
                tree_abonos.insert("", "end", values=(
                    row.fecha_registro, row.saldos, row.motivo, row.tipo
                ))
                total += float(row.saldos)
        actualizar_entry(entry_abonos, total)
        print(f"Total Abonos: {total:.0f}")
        actualizar_resumen(entry_deudas, entry_abonos, entry_actual)

    def actualizar_resumen(entry_deudas, entry_abonos, entry_actual):
        try:
            deuda_str = entry_deudas.get().replace(",", "").replace("$", "").strip()
            abonos_str = entry_abonos.get().replace(",", "").replace("$", "").strip()

            if not deuda_str or not abonos_str:
                print("[Aviso resumen] Uno de los campos está vacío, se omite la actualización.")
                return

            deuda = float(deuda_str)
            abonos = float(abonos_str)
            actual = deuda - abonos
            actualizar_entry(entry_actual, actual)
            print(f"Deuda Actual: {actual:.0f}")
        except Exception as e:
            print(f"[ERROR resumen] {e}")

    def on_cliente_dobleclick(event):
        selected = tree_clientes.focus()
        if not selected:
            return
        cedula, placa, nombre = tree_clientes.item(selected, "tags")
        cargar_abonos(nombre, placa, entry_total_abonos, entry_deuda_actual, entry_total_deudas)
        cargar_deudas(cedula, placa, entry_total_deudas, entry_deuda_actual, entry_total_abonos)

    def eliminar_deuda():
        selected = tree_deudas.selection()
        if not selected:
            messagebox.showwarning("Sin selección", "Selecciona una deuda para eliminar.")
            return

        confirm = messagebox.askyesno("Confirmar eliminación", "¿Estás seguro de eliminar esta deuda?")
        if not confirm:
            return

        item = tree_deudas.item(selected[0])
        id_deuda = item["values"][0]
        try:
            with engine.begin() as conn:
                conn.execute(tabla_otras_deudas.delete().where(tabla_otras_deudas.c.id == id_deuda))
            tree_deudas.delete(selected[0])
            item_cliente = tree_clientes.focus()
            if item_cliente:
                cedula, placa, nombre = tree_clientes.item(item_cliente, "tags")
                cargar_deudas(cedula, placa)
                cargar_abonos(nombre, placa)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo eliminar la deuda:\n{e}")

    def agregar_deuda():
        item_cliente = tree_clientes.focus()
        if not item_cliente:
            messagebox.showwarning("Selecciona un cliente", "Debes seleccionar un cliente primero.")
            return

        cedula, placa, nombre = tree_clientes.item(item_cliente, "tags")

        def guardar():
            descripcion = entry_desc.get().strip()
            valor = entry_valor.get().strip()
            if not descripcion or not valor:
                messagebox.showerror("Campos vacíos", "Todos los campos son obligatorios.")
                return
            try:
                valor_float = float(valor)
            except ValueError:
                messagebox.showerror("Valor inválido", "El valor debe ser un número.")
                return
            try:
                with engine.begin() as conn:
                    conn.execute(tabla_otras_deudas.insert().values(
                        cedula=cedula,
                        placa=placa,
                        fecha_deuda = datetime.today().date(),
                        descripcion=descripcion,
                        valor=valor_float
                    ))
                top.destroy()
                cargar_deudas(cedula, placa)
                cargar_abonos(nombre, placa)
                messagebox.showinfo("Éxito", "Deuda agregada correctamente.")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar la deuda.\n{e}")

        top = tk.Toplevel(ventana)
        top.title("Agregar Deuda")
        top.geometry("300x180")
        top.resizable(False, False)
        tk.Label(top, text=f"Cliente: {nombre}", font=("Arial", 10, "bold")).pack(pady=5)
        tk.Label(top, text="Descripción:").pack()
        entry_desc = tk.Entry(top, width=30, justify="center")
        entry_desc.pack(pady=3)
        tk.Label(top, text="Valor ($COP):").pack()
        entry_valor = tk.Entry(top, width=20, justify="center")
        entry_valor.pack(pady=3)
        tk.Button(top, text="Guardar", bg="#4CAF50", fg="white", command=guardar).pack(pady=10)

    def ver_consolidado():
        top = tk.Toplevel(ventana)
        top.title("Consolidado por Placa")
        top.geometry("600x500")
        top.configure(bg="white")

        cols = ("Placa", "Cliente", "Total Deudas", "Total Abonos", "Saldo Pendiente")
        tree_consolidado = ttk.Treeview(top, columns=cols, show="headings", height=20)

        for col in cols:
            tree_consolidado.heading(col, text=col)
            tree_consolidado.column(col, anchor="center")

        scroll = ttk.Scrollbar(top, orient="vertical", command=tree_consolidado.yview)
        tree_consolidado.configure(yscrollcommand=scroll.set)
        tree_consolidado.grid(row=0, column=0, sticky="nsew")
        scroll.grid(row=0, column=1, sticky="ns")

        top.grid_rowconfigure(0, weight=1)
        top.grid_columnconfigure(0, weight=1)

        with engine.begin() as conn:
            stmt_clientes = select(
                tabla_clientes.c.nombre,
                tabla_clientes.c.placa,
                tabla_clientes.c.cedula
            ).distinct()

            for nombre, placa, cedula in conn.execute(stmt_clientes):
                # Total deudas
                stmt_deudas = select(tabla_otras_deudas.c.valor).where(
                    tabla_otras_deudas.c.placa == placa,
                    tabla_otras_deudas.c.cedula == cedula
                )
                total_deudas = sum([float(row.valor) for row in conn.execute(stmt_deudas)])

                # Total abonos
                stmt_abonos = select(tabla_registros.c.saldos).where(
                    tabla_registros.c.placa == placa,
                    tabla_registros.c.nombre == nombre,
                    tabla_registros.c.saldos > 0,
                    tabla_registros.c.motivo.notin_(["N-a", "Multa"])
                )
                total_abonos = sum([float(row.saldos) for row in conn.execute(stmt_abonos)])

                saldo_pendiente = total_deudas - total_abonos

                tree_consolidado.insert("", "end", values=(
                    placa,
                    nombre,
                    f"{total_deudas:.0f}",
                    f"{total_abonos:.0f}",
                    f"{saldo_pendiente:.0f}"
                ))

    # --- INTERFAZ PRINCIPAL ---
    ventana = tk.Tk()
    ventana.title("Gestión de Deudas")
    ventana.geometry("1200x700")
    ventana.configure(bg="white")

    # --- Configuración general del grid ---
    ventana.grid_rowconfigure(0, weight=0)  # Fila superior: filtros, botones, resumen
    #ventana.grid_rowconfigure(1, weight=1)
    ventana.grid_rowconfigure(1, weight=0)  # Intermedio si lo usas
    ventana.grid_rowconfigure(2, weight=1)  # Área de expansión para clientes/abonos

    ventana.grid_columnconfigure(1, weight=1)
    ventana.grid_columnconfigure(0, weight=1)
    
    # Título principal
    label_titulo = tk.Label(ventana, text="GESTIÓN DE DEUDAS", font=("Arial", 14, "bold"), bg="white", fg="#2c3e50")
    label_titulo.grid(row=0, column=0, columnspan=2, pady=(10, 0))

    # Panel superior (izquierda): Filtros + Botones + Resumen
    frame_izq_superior = tk.Frame(ventana, bg="white")
    frame_izq_superior.grid(row=1, column=0, sticky="nsew", padx=10, pady=10)
    frame_izq_superior.grid_columnconfigure(0, weight=1)

    # --- Filtro por nombre y placa ---
    frame_filtro = tk.LabelFrame(frame_izq_superior, text="Filtros de búsqueda",
        font=("Arial", 10, "bold"), fg="black", bg="white",
        padx=10, pady=5, bd=1, relief="solid")
    frame_filtro.grid(row=0, column=0, sticky="ew", pady=5)

    tk.Label(frame_filtro, text="Filtrar por nombre:", font=("Arial", 10), bg="white").grid(row=0, column=0, sticky="nsew")
    entry_filtro_nombre = tk.Entry(frame_filtro, width=30)
    entry_filtro_nombre.grid(row=0, column=1, padx=(5, 15))

    tk.Label(frame_filtro, text="Filtrar por placa:", font=("Arial", 10), bg="white").grid(row=0, column=2, sticky="nsew")
    entry_filtro_placa = tk.Entry(frame_filtro, width=10)
    entry_filtro_placa.grid(row=0, column=3, padx=(5, 0))

    def actualizar_filtro_clientes(_event=None):
        cargar_clientes(entry_filtro_nombre.get(), entry_filtro_placa.get())

    entry_filtro_nombre.bind("<KeyRelease>", actualizar_filtro_clientes)
    entry_filtro_placa.bind("<KeyRelease>", actualizar_filtro_clientes)

    # --- Acciones + Resumen financiero ---
    frame_acciones_resumen = tk.LabelFrame(frame_izq_superior, text="Opciones y Resumen",
        font=("Arial", 10, "bold"), fg="black", bg="white",
        padx=10, pady=10, bd=1, relief="solid")
    frame_acciones_resumen.grid(row=1, column=0, sticky="nsew", pady=(10, 0), padx=0)
    frame_acciones_resumen.grid_columnconfigure(1, weight=1)

    # Botones
    btn_agregar = tk.Button(frame_acciones_resumen, text="Agregar Deuda", bg="#3498db", fg="white", command=agregar_deuda)
    btn_eliminar = tk.Button(frame_acciones_resumen, text="Eliminar Deuda", bg="#e74c3c", fg="white", command=eliminar_deuda)
    btn_consolidado = tk.Button(frame_acciones_resumen, text="📊 Ver Consolidado", bg="#38c500", fg="white", command=ver_consolidado)
    
    btn_agregar.grid(row=0, column=0, padx=5, pady=5, sticky="w")
    btn_eliminar.grid(row=0, column=1, padx=5, pady=5, sticky="w")
    btn_consolidado.grid(row=0, column=2, padx=5, pady=5, sticky="w")

    # Resumen
    tk.Label(frame_acciones_resumen, text="Total Deudas:", font=("Arial", 9, "bold"),
            bg="white", anchor="w").grid(row=1, column=0, sticky="w", pady=(10,0))
    entry_total_deudas = tk.Entry(frame_acciones_resumen, font=("Arial", 10, "bold"),
        fg="darkred", bg="white", bd=0, justify="right", state="readonly")
    entry_total_deudas.grid(row=1, column=1, sticky="e", pady=(10,0))

    tk.Label(frame_acciones_resumen, text="Total Abonos:", font=("Arial", 9, "bold"),
            bg="white", anchor="w").grid(row=2, column=0, sticky="w")
    entry_total_abonos = tk.Entry(frame_acciones_resumen, font=("Arial", 10, "bold"),
        fg="darkgreen", bg="white", bd=0, justify="right", state="readonly")
    entry_total_abonos.grid(row=2, column=1, sticky="e")

    tk.Label(frame_acciones_resumen, text="Deuda Actual:", font=("Arial", 9, "bold"),
            bg="white", anchor="w").grid(row=3, column=0, sticky="w")
    entry_deuda_actual = tk.Entry(frame_acciones_resumen, font=("Arial", 10, "bold"),
        fg="blue", bg="white", bd=0, justify="right", state="readonly")
    entry_deuda_actual.grid(row=3, column=1, sticky="e")

    # Panel derecho superior: Deudas
    frame_deudas = tk.LabelFrame(ventana, text="Deudas Registradas",
        font=("Arial", 10, "bold"), fg="black", bg="white",
        padx=10, pady=5, bd=1, relief="solid")
    frame_deudas.grid(row=1, column=1, sticky="nsew", padx=10, pady=10)
    frame_deudas.grid_rowconfigure(0, weight=1)
    frame_deudas.grid_columnconfigure(0, weight=1)

    cols_deudas = ("ID", "Fecha", "Descripción", "Valor")
    tree_deudas = ttk.Treeview(frame_deudas, columns=cols_deudas, show="headings")
    for col in cols_deudas:
        tree_deudas.heading(col, text=col)
        tree_deudas.column(col, anchor="center")

    scroll_deudas = ttk.Scrollbar(frame_deudas, orient="vertical", command=tree_deudas.yview)
    tree_deudas.configure(yscrollcommand=scroll_deudas.set)
    tree_deudas.grid(row=0, column=0, sticky="nsew")
    scroll_deudas.grid(row=0, column=1, sticky="ns")

    # Panel inferior izquierdo: Clientes
    frame_clientes = tk.LabelFrame(ventana, text="Lista de Clientes",
        font=("Arial", 10, "bold"), fg="black", bg="white",
        padx=10, pady=5, bd=1, relief="solid")
    frame_clientes.grid(row=2, column=0, sticky="nsew", padx=10, pady=10)
    frame_clientes.grid_rowconfigure(0, weight=1)
    frame_clientes.grid_columnconfigure(0, weight=1)

    cols_clientes = ("Nombre", "Placa")
    tree_clientes = ttk.Treeview(frame_clientes, columns=cols_clientes, show="headings")
    for col in cols_clientes:
        tree_clientes.heading(col, text=col)
        tree_clientes.column(col, anchor="center")

    scroll_clientes = ttk.Scrollbar(frame_clientes, orient="vertical", command=tree_clientes.yview)
    tree_clientes.configure(yscrollcommand=scroll_clientes.set)
    tree_clientes.grid(row=0, column=0, sticky="nsew")
    scroll_clientes.grid(row=0, column=1, sticky="ns")
    tree_clientes.bind("<Double-1>", on_cliente_dobleclick)

    # Panel inferior derecho: Abonos
    frame_abonos = tk.LabelFrame(ventana, text="Abonos Registrados",
        font=("Arial", 10, "bold"), fg="black", bg="white",
        padx=10, pady=5, bd=1, relief="solid")
    frame_abonos.grid(row=2, column=1, sticky="nsew", padx=10, pady=10)
    frame_abonos.grid_rowconfigure(0, weight=1)
    frame_abonos.grid_columnconfigure(0, weight=1)

    cols_abonos = ("Fecha", "Saldos", "Motivo", "Tipo")
    tree_abonos = ttk.Treeview(frame_abonos, columns=cols_abonos, show="headings")
    for col in cols_abonos:
        tree_abonos.heading(col, text=col)
        tree_abonos.column(col, anchor="center")

    scroll_abonos = ttk.Scrollbar(frame_abonos, orient="vertical", command=tree_abonos.yview)
    tree_abonos.configure(yscrollcommand=scroll_abonos.set)
    tree_abonos.grid(row=0, column=0, sticky="nsew")
    scroll_abonos.grid(row=0, column=1, sticky="ns")
    # Cargar inicial
    cargar_clientes()

#---------------- Función para iniciar la consulta de multas ----------
def iniciar_consulta_multas():
    # --- FUNCIONES ---
    def consultar_multas():
        fecha = date_entry.get_date()
        tree.delete(*tree.get_children())
        with engine.begin() as conn:
            stmt = select(tabla_registros).where(
                and_(
                    tabla_registros.c.motivo == "Multa",
                    tabla_registros.c.fecha_registro == fecha
                )
            ).order_by(tabla_registros.c.fecha_registro.desc())

            resultados = conn.execute(stmt).fetchall()
            if not resultados:
                messagebox.showinfo("Sin resultados", f"No hay multas para {fecha}")
                return

            for row in resultados:
                tree.insert("", "end", values=(
                    row.id, row.fecha_registro, row.cedula, row.nombre, row.placa,
                    row.valor, row.saldos, row.motivo, row.tipo, row.nombre_cuenta,
                    row.referencia, row.verificada
                ))
    # --- UI ---
    ventana = tk.Tk()
    ventana.title("Consultar Multas por Fecha")
    ventana.geometry("1300x500")
    ventana.configure(bg="white")

    # Filtro por fecha
    frame_filtro = tk.Frame(ventana, bg="white")
    frame_filtro.pack(pady=10)

    tk.Label(frame_filtro, text="Selecciona fecha:", bg="white", font=("Arial", 10, "bold")).pack(side="left")
    date_entry = DateEntry(frame_filtro, date_pattern="yyyy-mm-dd", width=12)
    date_entry.pack(side="left", padx=5)

    tk.Button(frame_filtro, text="Consultar", command=consultar_multas, bg="#3498db", fg="white").pack(side="left", padx=10)

    # Tabla resultados
    cols = ["ID", "Fecha", "Cédula", "Nombre", "Placa", "Valor", "Saldos", "Motivo", "Tipo", "Cuenta", "Referencia", "Verificada"]
    tree = ttk.Treeview(ventana, columns=cols, show="headings")
    for col in cols:
        tree.heading(col, text=col)
        tree.column(col, anchor="center")

    # Scrollbars
    scroll_y = ttk.Scrollbar(ventana, orient="vertical", command=tree.yview)
    scroll_x = ttk.Scrollbar(ventana, orient="horizontal", command=tree.xview)
    tree.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)

    tree.pack(fill="both", expand=True)
    scroll_y.pack(side="right", fill="y")
    scroll_x.pack(side="bottom", fill="x")

    ventana.mainloop()

# ---------- Función para lanzar el editor de registros ----------
def lanzar_editor_registros():
    metadata = MetaData()
    cuentas = Table("cuentas", metadata, autoload_with=engine)

    def get_nombres_cuenta():
        with Session(engine) as session:
            stmt = select(cuentas.c.nombre_cuenta).order_by(cuentas.c.nombre_cuenta)
            return [row[0] for row in session.execute(stmt).all()]

    root = tk.Toplevel()
    root.title("Consulta y Edición de Registros")
    root.geometry("1100x700")
    root.configure(padx=20, pady=20, bg="white")

    tipos_motivo = ["N-a", "Inicial", "Otras deudas", "Multa"]
    tipos_pago = ["Consignación", "Transfer Nequi", "Bancolombia", "Transfiya", "PTM", "Efectivo", "Ajuste P/P"]
    editable_cols = ["valor", "saldos", "fecha_registro", "motivo", "tipo", "referencia"]
    entries = {}
    cuentas_list = get_nombres_cuenta()
    selected_id = {"id": None}

    # --- Panel de edición ---
    edit_frame = tk.LabelFrame(root, text="Editar Registro", padx=10, pady=10, bg="white", font=("Arial", 11, "bold"))
    edit_frame.pack(pady=5, fill="x")

    field_grid = {
        "valor": (0, 0),
        "tipo": (1, 0),
        "saldos": (0, 1),
        "motivo": (1, 1),
        "fecha_registro": (0, 2),
        "referencia": (1, 2)
    }

    for field, (fila, columna) in field_grid.items():
        tk.Label(edit_frame, text=field + ":", anchor="e", bg="white").grid(row=fila, column=columna * 2, sticky="e", padx=5, pady=3)

        if field == "fecha_registro":
            e = DateEntry(edit_frame, width=23, date_pattern="yyyy-mm-dd")
        elif field == "motivo":
            e = ttk.Combobox(edit_frame, width=23, values=tipos_motivo, state="readonly")
        elif field == "tipo":
            e = ttk.Combobox(edit_frame, width=23, values=tipos_pago, state="readonly")
        else:
            e = tk.Entry(edit_frame, width=25)

        e.grid(row=fila, column=columna * 2 + 1, padx=5, pady=3)
        entries[field] = e

    # --- Botones de acción ---
    def limpiar_seleccion():
        for widget in entries.values():
            if isinstance(widget, DateEntry):
                widget.set_date(datetime.today())
            elif isinstance(widget, ttk.Combobox):
                widget.set('')
            else:
                widget.delete(0, tk.END)
        selected_id["id"] = None
        tree.selection_remove(tree.selection())

    def eliminar_registro():
        if not selected_id["id"]:
            messagebox.showwarning("Sin selección", "Selecciona un registro para eliminar.", parent=root)
            return
        confirm = messagebox.askyesno("Confirmar eliminación", "¿Estás seguro de que deseas eliminar este registro?", parent=root)
        if not confirm:
            return
        try:
            with Session(engine) as session:
                stmt = delete(registros).where(registros.c.id == int(selected_id["id"]))
                session.execute(stmt)
                session.commit()
            messagebox.showinfo("Eliminado", "Registro eliminado correctamente.", parent=root)
            limpiar_seleccion()
            cargar_registros()
        except SQLAlchemyError as e:
            messagebox.showerror("Error SQL", str(e), parent=root)

    botones_frame = tk.Frame(root, bg="white")
    botones_frame.pack(pady=5)

    def crear_boton(texto, comando, color="black", bg="#e0e0e0"):
        return tk.Button(botones_frame, text=texto, command=comando, font=("Arial", 10, "bold"),
                         bg=bg, fg=color, padx=12, pady=5, bd=1, relief="raised", cursor="hand2")

    crear_boton("Guardar Cambios", lambda: guardar_cambios(), "white", "#4CAF50").pack(side="left", padx=10)
    crear_boton("Eliminar Registro", eliminar_registro, "white", "#f44336").pack(side="left", padx=10)
    crear_boton("Limpiar Selección", limpiar_seleccion, "black", "#FFEB3B").pack(side="left", padx=10)

    # --- Filtro por fecha ---
    frame_top = tk.Frame(root, bg="white")
    frame_top.pack(pady=10)

    tk.Label(frame_top, text="Fecha sistema:", font=("Arial", 10), bg="white").pack(side="left", padx=(0, 5))
    fecha_selector = DateEntry(frame_top, date_pattern="yyyy-mm-dd")
    fecha_selector.pack(side="left", padx=5)
    tk.Button(frame_top, text="Buscar", command=lambda: cargar_registros(),
              font=("Arial", 10), bg="#2196F3", fg="white", padx=10).pack(side="left", padx=10)

    # --- Tabla de registros ---
    cols = [
        "id", "cedula", "nombre", "placa", "valor", "saldos",
        "motivo", "tipo", "fecha_registro", "nombre_cuenta", "referencia"
    ]

    tree_frame = tk.Frame(root)
    tree_frame.pack(pady=10, fill="both", expand=True)

    def ordenar_columna(treeview, col, reversa):
        datos = [(treeview.set(k, col), k) for k in treeview.get_children('')]

        # Intentar conversión numérica si aplica
        try:
            datos.sort(key=lambda t: float(t[0].replace("$", "").replace(",", "")), reverse=reversa)
        except ValueError:
            datos.sort(key=lambda t: t[0].lower(), reverse=reversa)

        for index, (val, k) in enumerate(datos):
            treeview.move(k, '', index)

        # Alternar orden en el siguiente clic
        treeview.heading(col, command=lambda: ordenar_columna(treeview, col, not reversa))


    # --- Tu parte original con la modificación incluida ---
    tree = ttk.Treeview(tree_frame, columns=cols, show="headings", height=12)

    for col in cols:
        tree.heading(col, text=col, command=lambda c=col: ordenar_columna(tree, c, False))  # Click ordenable
        tree.column(col, anchor="center", width=100)

    tree.pack(side="left", fill="both", expand=True)

    scrollbar = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
    scrollbar.pack(side="right", fill="y")
    tree.configure(yscrollcommand=scrollbar.set)


    # --- Funciones auxiliares ---
    def cargar_registros():
        tree.delete(*tree.get_children())
        fecha = fecha_selector.get_date()
        with Session(engine) as session:
            stmt = select(registros).where(registros.c.fecha_sistema == fecha)
            resultados = session.execute(stmt).mappings().all()

            if not resultados:
                messagebox.showinfo("Sin datos", f"No hay registros para {fecha.strftime('%Y-%m-%d')}", parent=root)
                return

            for fila in resultados:
                tree.insert("", "end", values=[fila[col] for col in cols])

    def on_double_click(event):
        item = tree.selection()
        if not item:
            return
        datos = tree.item(item[0], "values")
        selected_id["id"] = datos[0]

        for i, field in enumerate(cols[1:]):
            if field in editable_cols:
                widget = entries[field]
                value = datos[i + 1]
                if isinstance(widget, DateEntry):
                    try:
                        widget.set_date(datetime.strptime(value, "%Y-%m-%d").date())
                    except:
                        widget.set_date(datetime.today())
                elif isinstance(widget, ttk.Combobox):
                    widget.set(value)
                else:
                    widget.delete(0, tk.END)
                    widget.insert(0, value)

    tree.bind("<Double-1>", on_double_click)

    def guardar_cambios():
        if not selected_id["id"]:
            messagebox.showwarning("Sin selección", "Selecciona un registro para editar.", parent=root)
            return

        try:
            data = {f: entries[f].get() for f in editable_cols}

            valores = {
                "valor": float(data["valor"]),
                "saldos": float(data["saldos"]),
                "fecha_registro": entries["fecha_registro"].get_date(),
                "motivo": data["motivo"],
                "tipo": data["tipo"],
                "referencia": data["referencia"]
            }

            with Session(engine) as session:
                stmt = (
                    update(registros)
                    .where(registros.c.id == int(selected_id["id"]))
                    .values(**valores)
                )
                session.execute(stmt)
                session.commit()
                messagebox.showinfo("Éxito", "Registro actualizado correctamente.", parent=root)
                cargar_registros()

        except SQLAlchemyError as e:
            messagebox.showerror("Error SQL", str(e), parent=root)
        except ValueError:
            messagebox.showerror("Error", "Verifica que los campos 'valor', 'saldos' y 'fecha_registro' sean válidos.", parent=root)

    root.mainloop()


def normalizar_placa(placa):
    placa = placa.upper().strip()
    match = re.match(r'^([A-Z]{3}\d{2,3}[A-Z]?)', placa)
    return match.group(1) if match else placa

def formato_pesos(valor):
    return f"${int(round(valor)):,.0f}".replace(",", ".")

def lanzar_resumen_placas():
    ventana_resumen = tk.Toplevel()
    ventana_resumen.title("Resumen por placa (agrupadas)")
    ventana_resumen.geometry("600x500")
    ventana_resumen.configure(bg="white", padx=20, pady=20)

    # --- Tabla Treeview ---
    cols = ("placa_base", "total_valor", "total_saldos_inicial")
    tree = ttk.Treeview(ventana_resumen, columns=cols, show="headings", height=20)
    for col in cols:
        tree.heading(col, text=col)
        tree.column(col, anchor="center", width=180)
    tree.pack(fill="both", expand=True)

    # --- Consulta y carga ---
    def cargar_datos():
        tree.delete(*tree.get_children())
        placas_dict = {}

        with Session(engine) as session:
            stmt = select(
                registros.c.placa,
                registros.c.valor,
                registros.c.saldos,
                registros.c.motivo
            ).where(registros.c.placa != None)
            resultados = session.execute(stmt).all()

            for placa, valor, saldos, motivo in resultados:
                base = normalizar_placa(placa)

                if base not in placas_dict:
                    placas_dict[base] = {"valor": 0, "saldos_inicial": 0}

                placas_dict[base]["valor"] += float(valor or 0)
                if motivo == "Inicial":
                    placas_dict[base]["saldos_inicial"] += float(saldos or 0)

        # Ordenar por total valor descendente
        ordenado = sorted(placas_dict.items(), key=lambda x: x[1]["valor"], reverse=True)

        for placa, datos in ordenado:
            tree.insert("", "end", values=(
                placa,
                formato_pesos(datos["valor"]),
                formato_pesos(datos["saldos_inicial"])
            ))

    # --- Botón de carga ---
    btn_frame = tk.Frame(ventana_resumen, bg="white")
    btn_frame.pack(pady=10)
    tk.Button(
        btn_frame,
        text="Actualizar resumen",
        font=("Arial", 10, "bold"),
        bg="#4CAF50",
        fg="white",
        padx=10,
        pady=5,
        command=cargar_datos
    ).pack()

    cargar_datos()
